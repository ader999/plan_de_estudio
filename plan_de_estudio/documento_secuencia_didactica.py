from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from django.http import HttpResponse
from io import BytesIO
from .models import Silabo, AsignacionPlanEstudio, Guia
from django.utils import timezone
import datetime

def generar_documento_secuencia_didactica(silabos_agrupados, usuario, año_actual=None):
    """
    Genera un documento Word con la secuencia didáctica para todos los encuentros.
    
    Args:
        silabos_agrupados: Diccionario con silabos agrupados por código
        usuario: Nombre del usuario
        año_actual: Año actual (opcional)
        
    Returns:
        BytesIO: Buffer con el documento generado
    """
    # Crear un nuevo documento
    doc = Document()
    
    # Establecer márgenes de página
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Título del documento
    titulo = doc.add_heading('UNIVERSIDAD MARTÍN LUTERO', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitulo = doc.add_heading('SECUENCIA DIDÁCTICA', 1)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Si no se proporciona el año actual, usar el año actual
    if año_actual is None:
        año_actual = datetime.datetime.now().year
    
    # Procesar cada grupo de sílabos
    for codigo, silabos_grupo in silabos_agrupados.items():
        # Ordenar los sílabos por número de encuentro
        silabos_grupo.sort(key=lambda s: s.encuentros)
        
        # Información de la asignatura
        if silabos_grupo:
            primer_silabo = silabos_grupo[0]
            
            # Añadir sección de la asignatura
            doc.add_heading(f'Asignatura: {primer_silabo.asignacion_plan.plan_de_estudio.asignatura}', 2)
            doc.add_paragraph(f'Código: {primer_silabo.asignacion_plan.plan_de_estudio.codigo}')
            doc.add_paragraph(f'Carrera: {primer_silabo.asignacion_plan.plan_de_estudio.carrera.nombre}')
            doc.add_paragraph(f'Año académico: {primer_silabo.asignacion_plan.plan_de_estudio.año}')
            doc.add_paragraph(f'Trimestre: {primer_silabo.asignacion_plan.plan_de_estudio.trimestre}')
            doc.add_paragraph(f'Mediador: {usuario}')
            doc.add_paragraph(f'Año Lectivo: {año_actual}')
            
            # Procesar cada sílabo (encuentro) del grupo
            for silabo in silabos_grupo:
                # Añadir salto de página entre encuentros
                if silabo != silabos_grupo[0]:
                    doc.add_page_break()
                
                # Título del encuentro
                doc.add_heading(f'ENCUENTRO {silabo.encuentros}', 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Datos generales
                p = doc.add_heading('DATOS GENERALES', 3)
                
                # Tabla de datos generales
                tabla_datos = doc.add_table(rows=5, cols=2)
                tabla_datos.style = 'Table Grid'
                
                # Rellenar la tabla de datos generales
                tabla_datos.cell(0, 0).text = 'N° de guía:'
                tabla_datos.cell(0, 1).text = str(silabo.encuentros)
                tabla_datos.cell(1, 0).text = 'Fecha:'
                tabla_datos.cell(1, 1).text = silabo.fecha.strftime('%d/%m/%Y')
                tabla_datos.cell(2, 0).text = 'Sede:'
                tabla_datos.cell(2, 1).text = 'Jalapa'
                tabla_datos.cell(3, 0).text = 'Unidad:'
                tabla_datos.cell(3, 1).text = f'{silabo.unidad}. {silabo.nombre_de_la_unidad}'
                tabla_datos.cell(4, 0).text = 'Trimestre:'
                tabla_datos.cell(4, 1).text = silabo.asignacion_plan.plan_de_estudio.trimestre
                
                # Centrar título de unidad
                doc.add_heading('NOMBRE DE LA UNIDAD', 3).alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Tabla de nombre de unidad
                tabla_unidad = doc.add_table(rows=1, cols=1)
                tabla_unidad.style = 'Table Grid'
                tabla_unidad.cell(0, 0).text = f'{silabo.unidad}. {silabo.nombre_de_la_unidad}'
                set_cell_text_center(tabla_unidad.cell(0, 0))
                
                # Objetivos de la unidad
                doc.add_heading('OBJETIVOS DE LA UNIDAD', 3)
                
                tabla_objetivos = doc.add_table(rows=3, cols=2)
                tabla_objetivos.style = 'Table Grid'
                tabla_objetivos.cell(0, 0).text = 'Conceptual'
                tabla_objetivos.cell(0, 1).text = silabo.objetivo_conceptual
                tabla_objetivos.cell(1, 0).text = 'Procedimental'
                tabla_objetivos.cell(1, 1).text = silabo.objetivo_procedimental
                tabla_objetivos.cell(2, 0).text = 'Actitudinal'
                tabla_objetivos.cell(2, 1).text = silabo.objetivo_actitudinal
                
                # Establecer ancho de la primera columna
                for row in tabla_objetivos.rows:
                    row.cells[0].width = Inches(1.5)
                
                # Mediación Pedagógica
                doc.add_heading('MEDIACIÓN PEDAGÓGICA', 3).alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                tabla_mediacion = doc.add_table(rows=7, cols=6)
                tabla_mediacion.style = 'Table Grid'
                
                # Cabecera
                encabezados = ['Contenido temático', 'Fases', 'Forma organizativa', 
                               'Descripción de la forma organizativa', 'Recursos didácticos', 'Tiempo']
                
                for i, encabezado in enumerate(encabezados):
                    tabla_mediacion.cell(0, i).text = encabezado
                    set_cell_bold(tabla_mediacion.cell(0, i))
                
                # Contenido de la tabla
                # Contenido temático ocupa 6 filas
                tabla_mediacion.cell(1, 0).merge(tabla_mediacion.cell(6, 0))
                tabla_mediacion.cell(1, 0).text = silabo.contenido_tematico
                
                # Primer momento
                tabla_mediacion.cell(1, 1).merge(tabla_mediacion.cell(2, 1))
                set_cell_text_with_line_break(tabla_mediacion.cell(1, 1), "Primer momento", "(Entrada)")
                set_cell_bold(tabla_mediacion.cell(1, 1))
                
                tabla_mediacion.cell(1, 2).text = silabo.tipo_primer_momento
                tabla_mediacion.cell(1, 3).text = silabo.detalle_primer_momento
                tabla_mediacion.cell(1, 4).text = silabo.recursos_primer_momento
                tabla_mediacion.cell(1, 5).text = str(silabo.tiempo_primer_momento)
                
                # Espacio en blanco para la segunda fila del primer momento
                for i in range(2, 6):
                    tabla_mediacion.cell(2, i).text = "—"
                
                # Segundo momento
                tabla_mediacion.cell(3, 1).merge(tabla_mediacion.cell(4, 1))
                set_cell_text_with_line_break(tabla_mediacion.cell(3, 1), "Segundo momento", "(Elaboración)")
                set_cell_bold(tabla_mediacion.cell(3, 1))
                
                tabla_mediacion.cell(3, 2).text = f'Clases teóricas - {silabo.tipo_segundo_momento_claseteoria}'
                tabla_mediacion.cell(3, 3).text = silabo.clase_teorica
                tabla_mediacion.cell(3, 4).merge(tabla_mediacion.cell(4, 4))
                tabla_mediacion.cell(3, 4).text = "Computadora, guía, video"
                tabla_mediacion.cell(3, 5).text = "30"
                
                tabla_mediacion.cell(4, 2).text = "Clases prácticas - Evaluación formativa"
                tabla_mediacion.cell(4, 3).text = "Trabajo en grupos: ejercicios prácticos sobre arquitecturas."
                tabla_mediacion.cell(4, 5).text = "20"
                
                # Espacio en blanco para la fila 5
                for i in range(1, 6):
                    tabla_mediacion.cell(5, i).text = "—"
                
                # Tercer momento
                tabla_mediacion.cell(6, 1).text = "Tercer momento\n(Salida)"
                set_cell_bold(tabla_mediacion.cell(6, 1))
                tabla_mediacion.cell(6, 2).text = "Realimentación y orientación"
                tabla_mediacion.cell(6, 3).text = "Resumen del tema, aclaración de dudas, tareas para autoestudio."
                tabla_mediacion.cell(6, 4).text = "Pizarra, guía"
                tabla_mediacion.cell(6, 5).text = "20"
                
                # Ajustar anchos de columna
                set_column_width(tabla_mediacion, 0, 2.0)  # Contenido temático
                set_column_width(tabla_mediacion, 1, 1.5)  # Fases
                set_column_width(tabla_mediacion, 2, 1.5)  # Forma organizativa
                set_column_width(tabla_mediacion, 3, 3.0)  # Descripción
                set_column_width(tabla_mediacion, 4, 1.5)  # Recursos
                set_column_width(tabla_mediacion, 5, 1.0)  # Tiempo
                
                # Orientaciones de las actividades de aprendizaje
                doc.add_heading('ORIENTACIONES DE LAS ACTIVIDADES DE APRENDIZAJE', 3)
                
                tabla_actividades = doc.add_table(rows=2, cols=4)
                tabla_actividades.style = 'Table Grid'
                
                # Cabecera de tabla de actividades
                encabezados = ['Descripción de la actividad de Aprendizaje', 'Evaluación de la actividad de Aprendizaje', 
                               'Período', 'Tiempo en minutos']
                
                for i, encabezado in enumerate(encabezados):
                    tabla_actividades.cell(0, i).text = encabezado
                    set_cell_bold(tabla_actividades.cell(0, i))
                
                # Datos de la actividad
                actividad_aprendizaje = silabo.actividad_aprendizaje if silabo.actividad_aprendizaje else "Sin datos"
                tabla_actividades.cell(1, 0).text = actividad_aprendizaje
                
                # Evaluación
                texto_evaluacion = f"Tipo de Evaluación\n{silabo.tipo_evaluacion if silabo.tipo_evaluacion else 'Sin datos'}\n\n" \
                                  f"Instrumento de evaluación\n{silabo.instrumento_evaluacion if silabo.instrumento_evaluacion else 'Sin datos'}"
                tabla_actividades.cell(1, 1).text = texto_evaluacion
                
                # Período
                texto_periodo = f"{silabo.periodo_tiempo_programado if silabo.periodo_tiempo_programado else 'Sin datos'}\n\n" \
                                f"Técnica evaluación\n{silabo.tecnica_evaluacion if silabo.tecnica_evaluacion else 'Sin datos'}"
                tabla_actividades.cell(1, 2).text = texto_periodo
                
                # Tiempo
                texto_tiempo = f"{str(silabo.tiempo_minutos) if silabo.tiempo_minutos else 'Sin datos'}\n\n" \
                              f"Puntaje\n{str(silabo.puntaje) if silabo.puntaje else 'Sin datos'}"
                tabla_actividades.cell(1, 3).text = texto_tiempo
                
                # Añadir sección de Eje transversal
                doc.add_heading('EJE TRANSVERSAL', 3)
                
                tabla_eje = doc.add_table(rows=2, cols=2)
                tabla_eje.style = 'Table Grid'
                
                # Cabecera de tabla de eje transversal
                tabla_eje.cell(0, 0).text = 'Eje transversal'
                tabla_eje.cell(0, 1).text = 'Detalle eje transversal'
                set_cell_bold(tabla_eje.cell(0, 0))
                set_cell_bold(tabla_eje.cell(0, 1))
                
                # Datos del eje transversal
                tabla_eje.cell(1, 0).text = str(silabo.eje_transversal) if hasattr(silabo, 'eje_transversal') else "Sin datos"
                tabla_eje.cell(1, 1).text = silabo.detalle_eje_transversal if hasattr(silabo, 'detalle_eje_transversal') else "Sin datos"
                
                # Añadir GUÍA AUTODIDÁCTICA DE APRENDIZAJE si existe una guía asociada
                try:
                    # Intentar obtener la guía asociada al sílabo
                    guias = Guia.objects.filter(silabo=silabo)
                    
                    if guias.exists():
                        guia = guias.first()
                        
                        doc.add_heading('GUÍA AUTODIDÁCTICA DE APRENDIZAJE', 3).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Tabla de actividades de la guía
                        tabla_guia = doc.add_table(rows=5, cols=7)
                        tabla_guia.style = 'Table Grid'
                        
                        # Cabecera
                        encabezados_guia = ['N°', 'Nombre de la unidad', 'Objetivo de aprendizaje', 'Contenido temático', 
                                          'Actividad de aprendizaje', 'Evaluación', 'Tiempo']
                        
                        for i, encabezado in enumerate(encabezados_guia):
                            tabla_guia.cell(0, i).text = encabezado
                            set_cell_bold(tabla_guia.cell(0, i))
                        
                        # Actividad 1 (siempre presente)
                        tabla_guia.cell(1, 0).text = "1"
                        tabla_guia.cell(1, 1).text = guia.nombre_de_la_unidad
                        tabla_guia.cell(1, 2).text = f"{guia.tipo_objetivo_1}: {guia.objetivo_aprendizaje_1}"
                        tabla_guia.cell(1, 3).text = guia.contenido_tematico_1
                        tabla_guia.cell(1, 4).text = guia.actividad_aprendizaje_1
                        eval_text = f"Tipo: {guia.tipo_evaluacion_1}\nInstrumento: {guia.instrumento_evaluacion_1}\n"
                        eval_text += f"Recursos: {guia.recursos_didacticos_1}"
                        tabla_guia.cell(1, 5).text = eval_text
                        tiempo_text = f"{str(guia.tiempo_minutos_1)}\nPuntaje: {str(guia.puntaje_1)}"
                        tabla_guia.cell(1, 6).text = tiempo_text
                        
                        # Actividades opcionales (2-4)
                        # Actividad 2
                        if hasattr(guia, 'actividad_aprendizaje_2') and guia.actividad_aprendizaje_2:
                            tabla_guia.cell(2, 0).text = "2"
                            tabla_guia.cell(2, 1).text = guia.nombre_de_la_unidad
                            tabla_guia.cell(2, 2).text = f"{guia.tipo_objetivo_2}: {guia.objetivo_aprendizaje_2}"
                            tabla_guia.cell(2, 3).text = guia.contenido_tematico_2
                            tabla_guia.cell(2, 4).text = guia.actividad_aprendizaje_2
                            eval_text = f"Tipo: {guia.tipo_evaluacion_2}\nInstrumento: {guia.instrumento_evaluacion_2}\n"
                            eval_text += f"Recursos: {guia.recursos_didacticos_2}"
                            tabla_guia.cell(2, 5).text = eval_text
                            tiempo_text = f"{str(guia.tiempo_minutos_2)}\nPuntaje: {str(guia.puntaje_2)}"
                            tabla_guia.cell(2, 6).text = tiempo_text
                        else:
                            for i in range(7):
                                tabla_guia.cell(2, i).text = "—"
                        
                        # Actividad 3
                        if hasattr(guia, 'actividad_aprendizaje_3') and guia.actividad_aprendizaje_3:
                            tabla_guia.cell(3, 0).text = "3"
                            tabla_guia.cell(3, 1).text = guia.nombre_de_la_unidad
                            tabla_guia.cell(3, 2).text = f"{guia.tipo_objetivo_3}: {guia.objetivo_aprendizaje_3}"
                            tabla_guia.cell(3, 3).text = guia.contenido_tematico_3
                            tabla_guia.cell(3, 4).text = guia.actividad_aprendizaje_3
                            eval_text = f"Tipo: {guia.tipo_evaluacion_3}\nInstrumento: {guia.instrumento_evaluacion_3}\n"
                            eval_text += f"Recursos: {guia.recursos_didacticos_3}"
                            tabla_guia.cell(3, 5).text = eval_text
                            tiempo_text = f"{str(guia.tiempo_minutos_3)}\nPuntaje: {str(guia.puntaje_3)}"
                            tabla_guia.cell(3, 6).text = tiempo_text
                        else:
                            for i in range(7):
                                tabla_guia.cell(3, i).text = "—"
                        
                        # Actividad 4
                        if hasattr(guia, 'actividad_aprendizaje_4') and guia.actividad_aprendizaje_4:
                            tabla_guia.cell(4, 0).text = "4"
                            tabla_guia.cell(4, 1).text = guia.nombre_de_la_unidad
                            tabla_guia.cell(4, 2).text = f"{guia.tipo_objetivo_4}: {guia.objetivo_aprendizaje_4}"
                            tabla_guia.cell(4, 3).text = guia.contenido_tematico_4
                            tabla_guia.cell(4, 4).text = guia.actividad_aprendizaje_4
                            eval_text = f"Tipo: {guia.tipo_evaluacion_4}\nInstrumento: {guia.instrumento_evaluacion_4}\n"
                            eval_text += f"Recursos: {guia.recursos_didacticos_4}"
                            tabla_guia.cell(4, 5).text = eval_text
                            tiempo_text = f"{str(guia.tiempo_minutos_4)}\nPuntaje: {str(guia.puntaje_4)}"
                            tabla_guia.cell(4, 6).text = tiempo_text
                        else:
                            for i in range(7):
                                tabla_guia.cell(4, i).text = "—"
                        
                        # Calcular totales para guardar los datos
                        tiempo_total = 0
                        puntaje_total = 0
                        
                        # Sumar tiempo y puntaje de actividad 1 (obligatoria)
                        tiempo_total += guia.tiempo_minutos_1 if hasattr(guia, 'tiempo_minutos_1') else 0
                        puntaje_total += guia.puntaje_1 if hasattr(guia, 'puntaje_1') else 0
                        
                        # Sumar tiempo y puntaje de actividad 2 (opcional)
                        if hasattr(guia, 'tiempo_minutos_2') and guia.tiempo_minutos_2:
                            tiempo_total += guia.tiempo_minutos_2
                        if hasattr(guia, 'puntaje_2') and guia.puntaje_2:
                            puntaje_total += guia.puntaje_2
                        
                        # Sumar tiempo y puntaje de actividad 3 (opcional)
                        if hasattr(guia, 'tiempo_minutos_3') and guia.tiempo_minutos_3:
                            tiempo_total += guia.tiempo_minutos_3
                        if hasattr(guia, 'puntaje_3') and guia.puntaje_3:
                            puntaje_total += guia.puntaje_3
                        
                        # Sumar tiempo y puntaje de actividad 4 (opcional)
                        if hasattr(guia, 'tiempo_minutos_4') and guia.tiempo_minutos_4:
                            tiempo_total += guia.tiempo_minutos_4
                        if hasattr(guia, 'puntaje_4') and guia.puntaje_4:
                            puntaje_total += guia.puntaje_4
                            
                        # Resumen de horas y puntajes
                        doc.add_heading('RESUMEN DE HORAS Y PUNTAJES', 3)
                        
                        # Crear contenedor con dos columnas
                        tabla_container = doc.add_table(rows=1, cols=2)
                        tabla_container.style = 'Table Grid'
                        celda_izquierda = tabla_container.cell(0, 0)
                        celda_derecha = tabla_container.cell(0, 1)
                        
                        # Tabla izquierda para las horas
                        tabla_horas = celda_izquierda.add_table(rows=3, cols=2)
                        tabla_horas.style = 'Table Grid'
                        
                        # Cabeceras de horas
                        tabla_horas.cell(0, 0).text = 'Horas presenciales:'
                        tabla_horas.cell(1, 0).text = 'Horas de estudio independiente:'
                        tabla_horas.cell(2, 0).text = 'Total de horas:'
                        set_cell_bold(tabla_horas.cell(2, 0))
                        
                        # Calcular minutos presenciales
                        minutos_presenciales = 0
                        if hasattr(silabo, 'tiempo_primer_momento'):
                            minutos_presenciales += silabo.tiempo_primer_momento
                        if hasattr(silabo, 'tiempo_segundo_momento_teorica'):
                            minutos_presenciales += silabo.tiempo_segundo_momento_teorica
                        if hasattr(silabo, 'tiempo_segundo_momento_practica'):
                            minutos_presenciales += silabo.tiempo_segundo_momento_practica
                        if hasattr(silabo, 'tiempo_tercer_momento'):
                            minutos_presenciales += silabo.tiempo_tercer_momento
                        if hasattr(silabo, 'tiempo_minutos'):
                            minutos_presenciales += silabo.tiempo_minutos
                        
                        # Calcular horas presenciales
                        horas_presenciales = round(minutos_presenciales / 60, 2)
                        tabla_horas.cell(0, 1).text = f"{horas_presenciales}"
                        
                        # Calcular horas de estudio independiente
                        minutos_independientes = tiempo_total
                        horas_independientes = round(minutos_independientes / 60, 2)
                        tabla_horas.cell(1, 1).text = f"{horas_independientes}"
                        
                        # Total de horas
                        horas_totales = round((minutos_presenciales + minutos_independientes) / 60, 2)
                        tabla_horas.cell(2, 1).text = f"{horas_totales}"
                        
                        # Tabla derecha para los puntajes
                        tabla_puntajes = celda_derecha.add_table(rows=3, cols=2)
                        tabla_puntajes.style = 'Table Grid'
                        
                        # Cabeceras de puntajes
                        tabla_puntajes.cell(0, 0).text = 'Puntaje en actividad de aprendizaje presencial:'
                        tabla_puntajes.cell(1, 0).text = 'Puntaje en estudio independiente:'
                        tabla_puntajes.cell(2, 0).text = 'Puntaje Total:'
                        set_cell_bold(tabla_puntajes.cell(2, 0))
                        
                        # Puntaje presencial
                        puntaje_presencial = silabo.puntaje if hasattr(silabo, 'puntaje') else 0
                        tabla_puntajes.cell(0, 1).text = str(puntaje_presencial) if puntaje_presencial else "Sin datos"
                        
                        # Puntaje de estudio independiente
                        puntaje_independiente = puntaje_total
                        tabla_puntajes.cell(1, 1).text = str(puntaje_independiente) if puntaje_independiente else "Sin datos"
                        
                        # Puntaje total
                        puntaje_total_completo = puntaje_presencial + puntaje_independiente
                        tabla_puntajes.cell(2, 1).text = str(puntaje_total_completo) if puntaje_total_completo else "Sin datos"
                except Exception as e:
                    # Si hay un error al intentar agregar la guía, simplemente omitir esta sección
                    print(f"Error al intentar agregar la guía autodidáctica: {str(e)}")
    
    # Guardar documento en BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output

def set_cell_bold(cell):
    """Establece el texto de una celda como negrita."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

def set_cell_text_center(cell):
    """Centra el texto de una celda."""
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_cell_text_with_line_break(cell, text1, text2):
    """Establece el texto de una celda con salto de línea y formato."""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run1 = paragraph.add_run(text1)
    run1.bold = True
    
    paragraph.add_run("\n")
    paragraph.add_run(text2)

def set_column_width(table, column_index, width_inches):
    """Establece el ancho de una columna en pulgadas."""
    for cell in table.columns[column_index].cells:
        cell.width = Inches(width_inches)
