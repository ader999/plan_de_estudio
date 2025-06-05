from django.http import HttpResponse, JsonResponse
from django.contrib.auth.decorators import login_required
from openpyxl import load_workbook
from docx import Document
from .models import Guia, Silabo, AsignacionPlanEstudio
from .template_utils import get_excel_template

from django.contrib.admin.views.decorators import staff_member_required # Importar




# Helper function to safely write to cells, avoiding merged cell errors
def safe_write_cell(ws, row, col, value):
    """
    Safely writes a value to a cell, checking if it's a merged cell first.
    If it's a merged cell, writes to the master cell of the merge range.

    Args:
        ws: Worksheet object
        row: Row number (1-based)
        col: Column number (1-based)
        value: Value to write
    """
    # Get the cell
    cell = ws.cell(row=row, column=col)

    # Try to set the value directly
    try:
        cell.value = value
    except AttributeError:
        # If this fails, the cell might be part of a merged range
        print(f"Cell at row={row}, col={col} appears to be a merged cell. Trying to find the master cell.")

        # Find all merged ranges
        for merged_range in ws.merged_cells.ranges:
            # Check if the cell is in this range
            if cell.coordinate in merged_range:
                # Get the top-left cell (master cell) of the merged range
                master_cell = ws.cell(row=merged_range.min_row, column=merged_range.min_col)
                print(f"Writing to master cell at {master_cell.coordinate} instead")
                master_cell.value = value
                return

        # If we get here, the cell is not in any merged range, something else is wrong
        print(f"Failed to write to cell at row={row}, col={col}. Cell is not in any merged range.")

@login_required
def generar_excel_original(request):
    """
    Función para generar un archivo Excel basado en los datos de un sílabo.
    Ahora acepta tanto POST (id de sílabo) como GET (código de plan de estudio).
    """
    from django.http import JsonResponse
    silabo = None
    silabo_id = None
    codigo = None
    if request.method == 'POST':
        silabo_id = request.POST.get('codigoSilabo')
        print(f"ID de sílabo recibido: {silabo_id}")
        print(f"Usuario actual: {request.user.username}")
        if not silabo_id:
            return JsonResponse({'error': 'El ID de sílabo no se proporcionó.'}, status=400)
        try:
            print(f"Buscando sílabo con ID={silabo_id} y usuario={request.user.username}")
            silabo = Silabo.objects.get(id=silabo_id, asignacion_plan__usuario=request.user)
            print(f"Sílabo encontrado: ID={silabo.id}, Código={silabo.asignacion_plan.plan_de_estudio.codigo}")
        except Silabo.DoesNotExist:
            print(f"Error: No se encontró el sílabo con ID {silabo_id}")
            return JsonResponse({'error': f'No se encontró el sílabo con ID {silabo_id}'}, status=404)
    else:
        codigo = request.GET.get('codigo')
        print(f"Código recibido por GET: {codigo}")
        print(f"Usuario actual: {request.user.username}")
        if not codigo:
            return JsonResponse({'error': 'No se proporcionó el código de plan de estudio.'}, status=400)
        silabo = Silabo.objects.filter(
            asignacion_plan__plan_de_estudio__codigo=codigo,
            asignacion_plan__usuario=request.user
        ).order_by('encuentros').first()
        if not silabo:
            print(f"No se encontró sílabo para el usuario y código {codigo}")
            return JsonResponse({'error': f'No se encontró sílabo para el código {codigo}'}, status=404)
        print(f"Sílabo encontrado: ID={silabo.id}, Código={silabo.asignacion_plan.plan_de_estudio.codigo}")

        # Generamos el Excel para este sílabo
        print(f"Generando Excel para el sílabo ID={silabo.id}")

        # Carga la plantilla de Excel desde storage
        wb = get_excel_template('plantilla_original.xlsx')

        # Selecciona la segunda hoja de cálculo
        ws = wb.worksheets[1]  # Índice 1 para la segunda hoja

        # Datos generales del sílabo que solo se escriben una vez
        safe_write_cell(ws, 7, 4, silabo.asignacion_plan.usuario.first_name + " " + silabo.asignacion_plan.usuario.last_name)
        safe_write_cell(ws, 5, 10, silabo.asignacion_plan.plan_de_estudio.codigo)
        safe_write_cell(ws, 5, 11, silabo.asignacion_plan.plan_de_estudio.año)
        safe_write_cell(ws, 5, 12, silabo.asignacion_plan.plan_de_estudio.trimestre)
        safe_write_cell(ws, 5, 9, silabo.asignacion_plan.plan_de_estudio.asignatura.nombre)

        safe_write_cell(ws, 5, 6, silabo.asignacion_plan.plan_de_estudio.carrera.nombre)
        safe_write_cell(ws, 5, 3, silabo.asignacion_plan.plan_de_estudio.carrera.codigo)
        safe_write_cell(ws, 5, 4, silabo.asignacion_plan.plan_de_estudio.carrera.cine_2011)
        safe_write_cell(ws, 5, 5, silabo.asignacion_plan.plan_de_estudio.carrera.cine_2013)
        safe_write_cell(ws, 5, 7, silabo.asignacion_plan.plan_de_estudio.carrera.area_formacion)
        safe_write_cell(ws, 5, 8, silabo.asignacion_plan.plan_de_estudio.carrera.area_disiplinaria)

        safe_write_cell(ws, 7, 7, silabo.asignacion_plan.plan_de_estudio.horas_presenciales)
        safe_write_cell(ws, 7, 8, silabo.asignacion_plan.plan_de_estudio.horas_estudio_independiente)

        safe_write_cell(ws, 7, 10, silabo.asignacion_plan.plan_de_estudio.pr.nombre if silabo.asignacion_plan.plan_de_estudio.pr else "N/A")
        safe_write_cell(ws, 7, 11, silabo.asignacion_plan.plan_de_estudio.pc.nombre if silabo.asignacion_plan.plan_de_estudio.pc else "N/A")
        safe_write_cell(ws, 7, 12, silabo.asignacion_plan.plan_de_estudio.cr.nombre if silabo.asignacion_plan.plan_de_estudio.cr else "N/A")

        # Buscar todos los sílabos relacionados con la misma asignación de plan de estudio, ordenados por número de encuentro
        try:
                silabos_relacionados = Silabo.objects.filter(
                    asignacion_plan=silabo.asignacion_plan
                ).order_by('encuentros')

                print(f"Se encontraron {silabos_relacionados.count()} sílabos relacionados")

                # Definir la distancia en filas entre cada bloque de datos
                rows_per_block = 46
                rows_per_block_guia = 46

                # Iterar sobre cada sílabo y escribir sus datos con el offset apropiado
                for i, silabo_actual in enumerate(silabos_relacionados):
                    if i >= 11:  # Limitar a 11 sílabos máximo como solicitado
                        break

                    print(f"Procesando sílabo {i+1}: Encuentro {silabo_actual.encuentros}")

                    # Calcular el offset de filas para este sílabo
                    row_num = i * rows_per_block

                    # Escribir datos específicos del sílabo con el offset correspondiente
                    safe_write_cell(ws, 11+row_num, 3, silabo_actual.unidad+" " + silabo_actual.nombre_de_la_unidad)

                    safe_write_cell(ws, 12+row_num, 6, silabo_actual.objetivo_conceptual)
                    safe_write_cell(ws, 13+row_num, 6, silabo_actual.objetivo_procedimental)
                    safe_write_cell(ws, 18+row_num, 6, silabo_actual.objetivo_actitudinal)
                    safe_write_cell(ws, 11+row_num, 7, silabo_actual.contenido_tematico)

                    safe_write_cell(ws, 12+row_num, 8, silabo_actual.tipo_primer_momento)
                    safe_write_cell(ws, 12+row_num, 9, silabo_actual.detalle_primer_momento)
                    safe_write_cell(ws, 12+row_num, 10, silabo_actual.tiempo_primer_momento)
                    safe_write_cell(ws, 12+row_num, 11, silabo_actual.recursos_primer_momento)

                    safe_write_cell(ws, 15+row_num, 8, silabo_actual.tipo_segundo_momento_claseteoria)
                    safe_write_cell(ws, 14+row_num, 9, silabo_actual.clase_teorica)
                    safe_write_cell(ws, 17+row_num, 8, silabo_actual.tipo_segundo_momento_practica)
                    safe_write_cell(ws, 16+row_num, 9, silabo_actual.clase_practica)
                    safe_write_cell(ws, 14+row_num, 10, silabo_actual.tiempo_segundo_momento_teorica)
                    safe_write_cell(ws, 16+row_num, 10, silabo_actual.tiempo_segundo_momento_practica)
                    safe_write_cell(ws, 14+row_num, 11, silabo_actual.recursos_segundo_momento)

                    safe_write_cell(ws, 19+row_num, 8, ", ".join(silabo_actual.tipo_tercer_momento[:2]))
                    safe_write_cell(ws, 19+row_num, 9, silabo_actual.detalle_tercer_momento)
                    safe_write_cell(ws, 19+row_num, 10, silabo_actual.tiempo_tercer_momento)
                    safe_write_cell(ws, 19+row_num, 11, silabo_actual.recursos_tercer_momento)

                    safe_write_cell(ws, 11+row_num, 12, ", ".join(silabo_actual.eje_transversal[:2]))
                    safe_write_cell(ws, 12+row_num, 12, silabo_actual.detalle_eje_transversal)

                    #evaluacion dinamica
                    safe_write_cell(ws, 12+row_num, 14, silabo_actual.actividad_aprendizaje)
                    safe_write_cell(ws, 13+row_num, 14, ", ".join(silabo_actual.tecnica_evaluacion[:2]))
                    safe_write_cell(ws, 14+row_num, 14, ", ".join(silabo_actual.tipo_evaluacion[:2]))
                    safe_write_cell(ws, 15+row_num, 14, silabo_actual.periodo_tiempo_programado)
                    safe_write_cell(ws, 16+row_num, 14, silabo_actual.tiempo_minutos)
                    safe_write_cell(ws, 17+row_num, 14, ", ".join(silabo_actual.agente_evaluador[:2]))
                    safe_write_cell(ws, 18+row_num, 14, silabo_actual.instrumento_evaluacion)
                    safe_write_cell(ws, 19+row_num, 14, silabo_actual.criterios_evaluacion)
                    safe_write_cell(ws, 20+row_num, 14, silabo_actual.puntaje)

                    # Datos de la guía de estudio independiente (si existe)
                    if hasattr(silabo_actual, 'guia') and silabo_actual.guia is not None:
                        guia = silabo_actual.guia  # Accedemos a la instancia de guía
                        try:
                            if i == 1:
                                row_num_guia = 0

                            row_num_guia = i * rows_per_block_guia

                            safe_write_cell(ws, 23+row_num_guia, 2, guia.fecha if hasattr(guia, 'fecha') else "")
                            safe_write_cell(ws, 23+row_num_guia, 3, f"{guia.unidad} {guia.nombre_de_la_unidad}" if hasattr(guia, 'unidad') and hasattr(guia, 'nombre_de_la_unidad') else "")

                            safe_write_cell(ws, 23+row_num_guia, 4, guia.tipo_objetivo_1 if hasattr(guia, 'tipo_objetivo_1') else "")
                            safe_write_cell(ws, 24+row_num_guia, 6, guia.objetivo_aprendizaje_1 if hasattr(guia, 'objetivo_aprendizaje_1') else "")
                            safe_write_cell(ws, 24+row_num_guia, 7, guia.contenido_tematico_1 if hasattr(guia, 'contenido_tematico_1') else "")
                            safe_write_cell(ws, 23+row_num_guia, 9, guia.actividad_aprendizaje_1)
                            safe_write_cell(ws, 25+row_num_guia, 9, guia.tecnica_evaluacion_1)
                            safe_write_cell(ws, 26+row_num_guia, 9, guia.tipo_evaluacion_1)
                            safe_write_cell(ws, 27+row_num_guia, 9, guia.instrumento_evaluacion_1)
                            safe_write_cell(ws, 28+row_num_guia, 9, guia.criterios_evaluacion_1)
                            safe_write_cell(ws, 29+row_num_guia, 9, ", ".join(guia.agente_evaluador_1[:2]))
                            safe_write_cell(ws, 23+row_num_guia, 10, guia.tiempo_minutos_1)
                            safe_write_cell(ws, 23+row_num_guia, 11, guia.recursos_didacticos_1)
                            safe_write_cell(ws, 23+row_num_guia, 13, guia.periodo_tiempo_programado_1)
                            safe_write_cell(ws, 24+row_num_guia, 13, guia.puntaje_1)
                            safe_write_cell(ws, 25+row_num_guia, 14, guia.fecha_entrega_1)

                            safe_write_cell(ws, 31+row_num_guia, 4, guia.tipo_objetivo_2 if hasattr(guia, 'tipo_objetivo_2') else "")
                            safe_write_cell(ws, 31+row_num_guia, 6, guia.objetivo_aprendizaje_2 if hasattr(guia, 'objetivo_aprendizaje_2') else "")
                            safe_write_cell(ws, 31+row_num_guia, 7, guia.contenido_tematico_2 if hasattr(guia, 'contenido_tematico_2') else "")
                            safe_write_cell(ws, 31+row_num_guia, 9, guia.actividad_aprendizaje_2)
                            safe_write_cell(ws, 33+row_num_guia, 9, guia.tecnica_evaluacion_2)
                            safe_write_cell(ws, 34+row_num_guia, 9, guia.tipo_evaluacion_2)
                            safe_write_cell(ws, 35+row_num_guia, 9, guia.instrumento_evaluacion_2)
                            safe_write_cell(ws, 36+row_num_guia, 9, guia.criterios_evaluacion_2)
                            safe_write_cell(ws, 37+row_num_guia, 9, ", ".join(guia.agente_evaluador_2[:2]))
                            safe_write_cell(ws, 31+row_num_guia, 10, guia.tiempo_minutos_2)
                            safe_write_cell(ws, 31+row_num_guia, 11, guia.recursos_didacticos_2)
                            safe_write_cell(ws, 31+row_num_guia, 13, guia.periodo_tiempo_programado_2)
                            safe_write_cell(ws, 32+row_num_guia, 13, guia.puntaje_2)
                            safe_write_cell(ws, 33+row_num_guia, 14, guia.fecha_entrega_2)

                            safe_write_cell(ws, 39+row_num_guia, 4, guia.tipo_objetivo_3 if hasattr(guia, 'tipo_objetivo_3') else "")
                            safe_write_cell(ws, 39+row_num_guia, 6, guia.objetivo_aprendizaje_3 if hasattr(guia, 'objetivo_aprendizaje_3') else "")
                            safe_write_cell(ws, 39+row_num_guia, 7, guia.contenido_tematico_3 if hasattr(guia, 'contenido_tematico_3') else "")
                            safe_write_cell(ws, 39+row_num_guia, 9, guia.actividad_aprendizaje_3)
                            safe_write_cell(ws, 41+row_num_guia, 9, guia.tecnica_evaluacion_3)
                            safe_write_cell(ws, 42+row_num_guia, 9, guia.tipo_evaluacion_3)
                            safe_write_cell(ws, 43+row_num_guia, 9, guia.instrumento_evaluacion_3)
                            safe_write_cell(ws, 44+row_num_guia, 9, guia.criterios_evaluacion_3)
                            safe_write_cell(ws, 45+row_num_guia, 9, ", ".join(guia.agente_evaluador_3[:2]))
                            safe_write_cell(ws, 39+row_num_guia, 10, guia.tiempo_minutos_3)
                            safe_write_cell(ws, 39+row_num_guia, 11, guia.recursos_didacticos_3)
                            safe_write_cell(ws, 39+row_num_guia, 13, guia.periodo_tiempo_programado_3)
                            safe_write_cell(ws, 40+row_num_guia, 13, guia.puntaje_3)
                            safe_write_cell(ws, 41+row_num_guia, 14, guia.fecha_entrega_3)

                            safe_write_cell(ws, 47+row_num_guia, 4, guia.tipo_objetivo_4 if hasattr(guia, 'tipo_objetivo_1') else "")
                            safe_write_cell(ws, 47+row_num_guia, 6, guia.objetivo_aprendizaje_4 if hasattr(guia, 'objetivo_aprendizaje_1') else "")
                            safe_write_cell(ws, 47+row_num_guia, 7, guia.contenido_tematico_4 if hasattr(guia, 'contenido_tematico_1') else "")
                            safe_write_cell(ws, 47+row_num_guia, 9, guia.actividad_aprendizaje_4)
                            safe_write_cell(ws, 49+row_num_guia, 9, guia.tecnica_evaluacion_4)
                            safe_write_cell(ws, 50+row_num_guia, 9, guia.tipo_evaluacion_4)
                            safe_write_cell(ws, 51+row_num_guia, 9, guia.instrumento_evaluacion_4)
                            safe_write_cell(ws, 52+row_num_guia, 9, guia.criterios_evaluacion_4)
                            safe_write_cell(ws, 53+row_num_guia, 9, ", ".join(guia.agente_evaluador_4[:2]))
                            safe_write_cell(ws, 47+row_num_guia, 10, guia.tiempo_minutos_4)
                            safe_write_cell(ws, 47+row_num_guia, 11, guia.recursos_didacticos_4)
                            safe_write_cell(ws, 47+row_num_guia, 13, guia.periodo_tiempo_programado_4)
                            safe_write_cell(ws, 48+row_num_guia, 13, guia.puntaje_4)
                            safe_write_cell(ws, 49+row_num_guia, 14, guia.fecha_entrega_4)




                            # Podemos agregar más campos si es necesario
                            print(f"Guía procesada correctamente para el sílabo {silabo_actual.encuentros}")
                        except Exception as e:
                            print(f"Error al procesar la guía: {str(e)}")
                            # Continuamos con el siguiente sílabo sin interrumpir el proceso
                print("Generando respuesta con el archivo Excel...")
                # Guarda el archivo Excel en memoria
                response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                response['Content-Disposition'] = f'attachment; filename=silabo_{silabo.asignacion_plan.usuario.username+ " "+silabo.asignacion_plan.plan_de_estudio.asignatura.nombre}.xlsx'
                wb.save(response)
                print("Archivo Excel generado correctamente")
                return response

        except Exception as e:
            print(f"Error específico durante la generación del Excel: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return JsonResponse({'error': f'Error al generar el Excel: {str(e)}'}, status=500)

        except Silabo.DoesNotExist:
            print(f"Error: No se encontró el sílabo con ID {silabo_id}")
            return JsonResponse({'error': f'No se encontró el sílabo con ID {silabo_id}'}, status=404)
        except Exception as e:
            print(f"Error general al generar el Excel: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return JsonResponse({'error': f'Error al generar el Excel: {str(e)}'}, status=500)

    # Si no se proporcionó un ID de sílabo o no se encontró, redirige a la página principal
    from django.shortcuts import redirect
    return redirect('plan_de_estudio')


@login_required
def generar_docx(request):
    if request.method == 'POST':
        silabo_id = request.POST.get('codigoSilabo', '')
        print(f"ID de sílabo recibido para Word: {silabo_id}")
        print(f"Usuario actual: {request.user.username}")

        if not silabo_id:
            return JsonResponse({'error': 'El ID de sílabo no se proporcionó.'}, status=400)

        try:
            # Si se proporcionó un ID de sílabo, busca el sílabo directamente
            print(f"Buscando sílabo con ID={silabo_id} y usuario={request.user.username}")
            silabo = Silabo.objects.get(id=silabo_id, asignacion_plan__usuario=request.user)
            print(f"Sílabo encontrado: ID={silabo.id}, Código={silabo.codigo}")

            # Generamos el Word para este sílabo
            print(f"Generando Word para el sílabo ID={silabo.id}")

            # Crear un nuevo documento de Word
            doc = Document()

            # Añadir título
            doc.add_heading(f'Sílabo de {silabo.asignacion_plan.plan_de_estudio.asignatura.nombre}', 0)

            # Información general
            doc.add_heading('Información General', level=1)
            doc.add_paragraph(f'Profesor: {silabo.asignacion_plan.usuario.username}')
            doc.add_paragraph(f'Carrera: {silabo.asignacion_plan.plan_de_estudio.carrera.nombre}')
            doc.add_paragraph(f'Año: {silabo.asignacion_plan.plan_de_estudio.año}')
            doc.add_paragraph(f'Trimestre: {silabo.asignacion_plan.plan_de_estudio.trimestre}')
            doc.add_paragraph(f'Código: {silabo.codigo}')
            doc.add_paragraph(f'Encuentro: {silabo.encuentros}')
            doc.add_paragraph(f'Fecha: {silabo.fecha}')

            # Objetivos
            doc.add_heading('Objetivos', level=1)
            doc.add_paragraph(f'Conceptual: {silabo.objetivo_conceptual}')
            doc.add_paragraph(f'Procedimental: {silabo.objetivo_procedimental}')
            doc.add_paragraph(f'Actitudinal: {silabo.objetivo_actitudinal}')

            # Contenido
            doc.add_heading('Contenido Temático', level=1)
            doc.add_paragraph(f'Unidad: {silabo.unidad}')
            doc.add_paragraph(f'Nombre de la Unidad: {silabo.nombre_de_la_unidad}')
            doc.add_paragraph(f'Contenido: {silabo.contenido_tematico}')

            # Momentos Didácticos
            doc.add_heading('Momentos Didácticos', level=1)

            table = doc.add_table(rows=12, cols=2)
            table.style = 'Table Grid'

            cells = table.rows[0].cells
            cells[0].text = 'Campo'
            cells[1].text = 'Descripción'

            cells = table.rows[1].cells
            cells[0].text = 'Primer Momento'
            cells[1].text = silabo.detalle_primer_momento

            cells = table.rows[2].cells
            cells[0].text = 'Clase Teórica'
            cells[1].text = silabo.clase_teorica

            cells = table.rows[3].cells
            cells[0].text = 'Clase Práctica'
            cells[1].text = silabo.clase_practica

            cells = table.rows[4].cells
            cells[0].text = 'Tercer Momento'
            cells[1].text = silabo.detalle_tercer_momento

            cells = table.rows[5].cells
            cells[0].text = 'Contenido Temático'
            cells[1].text = silabo.contenido_tematico

            # Agregar información de la guía si existe
            if hasattr(silabo, 'guia') and silabo.guia is not None:
                doc.add_heading('Guía de Estudio Independiente', level=1)
                doc.add_paragraph(f'Unidad: {silabo.guia.unidad}')
                doc.add_paragraph(f'Contenido Temático: {silabo.guia.contenido_tematico}')
                doc.add_paragraph(f'Objetivo Conceptual: {silabo.guia.objetivo_conceptual}')
                doc.add_paragraph(f'Objetivo Procedimental: {silabo.guia.objetivo_procedimental}')
                doc.add_paragraph(f'Objetivo Actitudinal: {silabo.guia.objetivo_actitudinal}')
                doc.add_paragraph(f'Tiempo en Minutos: {silabo.guia.tiempo_minutos}')
                doc.add_paragraph(f'Puntaje: {silabo.guia.puntaje}')
                doc.add_paragraph(f'Criterios de Evaluación: {silabo.guia.criterios_evaluacion}')

            # Guardar el documento en memoria
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename=silabo_{silabo.codigo}.docx'

            # Guardar el documento en el objeto response
            doc.save(response)

            return response

        except Silabo.DoesNotExist:
            return JsonResponse({'error': f'No se encontró el sílabo con ID {silabo_id} para este usuario.'}, status=404)
        except Exception as e:
            return JsonResponse({'error': f'Error al generar el documento Word: {str(e)}'}, status=500)

    # Si no es una petición POST, redirige a la página principal
    from django.shortcuts import redirect
    return redirect('plan_de_estudio')




@staff_member_required # Asegurar que solo el personal del admin pueda acceder
def generar_excel_admin(request, asignacion_id):
    """
    Genera un archivo Excel para todos los sílabos y guías
    asociados a una AsignacionPlanEstudio específica.
    Activado desde el panel de administración.
    """
    print(f"Iniciando exportación de Excel desde admin para Asignacion ID: {asignacion_id}")
    try:
        # Busca la AsignacionPlanEstudio específica
        asignacion_obj = AsignacionPlanEstudio.objects.get(id=asignacion_id)
        print(f"Asignación encontrada: Usuario={asignacion_obj.usuario.username}, Plan={asignacion_obj.plan_de_estudio}")

        # Buscar todos los sílabos relacionados con esta asignación, ordenados por número de encuentro
        silabos_relacionados = Silabo.objects.filter(
            asignacion_plan=asignacion_obj
        ).order_by('encuentros').prefetch_related('guia') # prefetch_related para eficiencia

        if not silabos_relacionados.exists():
            # Podrías retornar un mensaje o un excel vacío si prefieres
            print(f"No se encontraron sílabos para la Asignación ID: {asignacion_id}")
            return HttpResponse("No se encontraron sílabos para esta asignación.", status=404)

        print(f"Se encontraron {silabos_relacionados.count()} sílabos relacionados para la Asignación ID: {asignacion_id}")

        # --- Lógica de generación de Excel (adaptada de generar_excel_original) ---

        # Carga la plantilla de Excel desde storage
        wb = get_excel_template('plantilla_original.xlsx')

        # Selecciona la segunda hoja de cálculo (o la que necesites)
        ws = wb.worksheets[1]  # Índice 1 para la segunda hoja

        # --- Escribir datos generales (usando asignacion_obj) ---
        plan = asignacion_obj.plan_de_estudio
        usuario = asignacion_obj.usuario
        carrera = plan.carrera
        asignatura = plan.asignatura

        safe_write_cell(ws, 7, 4, f"{usuario.first_name} {usuario.last_name}")
        safe_write_cell(ws, 5, 10, plan.codigo)
        safe_write_cell(ws, 5, 11, plan.año)
        safe_write_cell(ws, 5, 12, plan.trimestre)
        safe_write_cell(ws, 5, 9, asignatura.nombre)

        safe_write_cell(ws, 5, 6, carrera.nombre)
        safe_write_cell(ws, 5, 3, carrera.codigo)
        safe_write_cell(ws, 5, 4, carrera.cine_2011)
        safe_write_cell(ws, 5, 5, carrera.cine_2013)
        safe_write_cell(ws, 5, 7, carrera.area_formacion)
        safe_write_cell(ws, 5, 8, carrera.area_disiplinaria)

        safe_write_cell(ws, 7, 7, plan.horas_presenciales)
        safe_write_cell(ws, 7, 8, plan.horas_estudio_independiente)

        safe_write_cell(ws, 7, 10, plan.pr.nombre if plan.pr else "N/A")
        safe_write_cell(ws, 7, 11, plan.pc.nombre if plan.pc else "N/A")
        safe_write_cell(ws, 7, 12, plan.cr.nombre if plan.cr else "N/A")

        # Definir la distancia en filas entre cada bloque de datos
        rows_per_block = 46
        rows_per_block_guia = 46

        # --- Iterar sobre cada sílabo y escribir sus datos ---
        for i, silabo_actual in enumerate(silabos_relacionados):
            # if i >= 11: # Mantener el límite si es necesario
            #     break

            print(f"Procesando sílabo {i+1} (ID: {silabo_actual.id}): Encuentro {silabo_actual.encuentros}")

            # Calcular el offset de filas para este sílabo
            row_num = i * rows_per_block

            # --- Escribir datos específicos del sílabo (igual que en generar_excel_original) ---
            safe_write_cell(ws, 11+row_num, 3, f"{silabo_actual.unidad} {silabo_actual.nombre_de_la_unidad}")
            safe_write_cell(ws, 12+row_num, 6, silabo_actual.objetivo_conceptual)
            safe_write_cell(ws, 13+row_num, 6, silabo_actual.objetivo_procedimental)
            safe_write_cell(ws, 18+row_num, 6, silabo_actual.objetivo_actitudinal)
            safe_write_cell(ws, 11+row_num, 7, silabo_actual.contenido_tematico)
            safe_write_cell(ws, 12+row_num, 8, silabo_actual.tipo_primer_momento)
            safe_write_cell(ws, 12+row_num, 9, silabo_actual.detalle_primer_momento)
            safe_write_cell(ws, 12+row_num, 10, silabo_actual.tiempo_primer_momento)
            safe_write_cell(ws, 12+row_num, 11, silabo_actual.recursos_primer_momento)
            safe_write_cell(ws, 15+row_num, 8, silabo_actual.tipo_segundo_momento_claseteoria)
            safe_write_cell(ws, 14+row_num, 9, silabo_actual.clase_teorica)
            safe_write_cell(ws, 17+row_num, 8, silabo_actual.tipo_segundo_momento_practica)
            safe_write_cell(ws, 16+row_num, 9, silabo_actual.clase_practica)
            safe_write_cell(ws, 14+row_num, 10, silabo_actual.tiempo_segundo_momento_teorica)
            safe_write_cell(ws, 16+row_num, 10, silabo_actual.tiempo_segundo_momento_practica)
            safe_write_cell(ws, 14+row_num, 11, silabo_actual.recursos_segundo_momento)
            safe_write_cell(ws, 19+row_num, 8, ", ".join(silabo_actual.tipo_tercer_momento[:2] if silabo_actual.tipo_tercer_momento else []))
            safe_write_cell(ws, 19+row_num, 9, silabo_actual.detalle_tercer_momento)
            safe_write_cell(ws, 19+row_num, 10, silabo_actual.tiempo_tercer_momento)
            safe_write_cell(ws, 19+row_num, 11, silabo_actual.recursos_tercer_momento)
            safe_write_cell(ws, 11+row_num, 12, ", ".join(silabo_actual.eje_transversal[:2] if silabo_actual.eje_transversal else []))
            safe_write_cell(ws, 12+row_num, 12, silabo_actual.detalle_eje_transversal)
            safe_write_cell(ws, 12+row_num, 14, silabo_actual.actividad_aprendizaje)
            safe_write_cell(ws, 13+row_num, 14, ", ".join(silabo_actual.tecnica_evaluacion[:2] if silabo_actual.tecnica_evaluacion else []))
            safe_write_cell(ws, 14+row_num, 14, ", ".join(silabo_actual.tipo_evaluacion[:2] if silabo_actual.tipo_evaluacion else []))
            safe_write_cell(ws, 15+row_num, 14, silabo_actual.periodo_tiempo_programado)
            safe_write_cell(ws, 16+row_num, 14, silabo_actual.tiempo_minutos)
            safe_write_cell(ws, 17+row_num, 14, ", ".join(silabo_actual.agente_evaluador[:2] if silabo_actual.agente_evaluador else []))
            safe_write_cell(ws, 18+row_num, 14, silabo_actual.instrumento_evaluacion)
            safe_write_cell(ws, 19+row_num, 14, silabo_actual.criterios_evaluacion)
            safe_write_cell(ws, 20+row_num, 14, silabo_actual.puntaje)

            # --- Datos de la guía de estudio independiente (si existe) ---
            if hasattr(silabo_actual, 'guia') and silabo_actual.guia is not None:
                guia = silabo_actual.guia
                try:
                    row_num_guia = i * rows_per_block_guia # Calcular offset para la guía

                    # --- Escribir datos de la guía (igual que en generar_excel_original) ---
                    safe_write_cell(ws, 23+row_num_guia, 2, guia.fecha if hasattr(guia, 'fecha') else "")
                    safe_write_cell(ws, 23+row_num_guia, 3, f"{guia.unidad} {guia.nombre_de_la_unidad}" if hasattr(guia, 'unidad') and hasattr(guia, 'nombre_de_la_unidad') else "")
                    safe_write_cell(ws, 23+row_num_guia, 4, guia.tipo_objetivo_1 if hasattr(guia, 'tipo_objetivo_1') else "")
                    safe_write_cell(ws, 24+row_num_guia, 6, guia.objetivo_aprendizaje_1 if hasattr(guia, 'objetivo_aprendizaje_1') else "")
                    safe_write_cell(ws, 24+row_num_guia, 7, guia.contenido_tematico_1 if hasattr(guia, 'contenido_tematico_1') else "")
                    safe_write_cell(ws, 23+row_num_guia, 9, guia.actividad_aprendizaje_1)
                    safe_write_cell(ws, 25+row_num_guia, 9, guia.tecnica_evaluacion_1)
                    safe_write_cell(ws, 26+row_num_guia, 9, guia.tipo_evaluacion_1)
                    safe_write_cell(ws, 27+row_num_guia, 9, guia.instrumento_evaluacion_1)
                    safe_write_cell(ws, 28+row_num_guia, 9, guia.criterios_evaluacion_1)
                    safe_write_cell(ws, 29+row_num_guia, 9, ", ".join(guia.agente_evaluador_1[:2] if guia.agente_evaluador_1 else []))
                    safe_write_cell(ws, 23+row_num_guia, 10, guia.tiempo_minutos_1)
                    safe_write_cell(ws, 23+row_num_guia, 11, guia.recursos_didacticos_1)
                    safe_write_cell(ws, 23+row_num_guia, 13, guia.periodo_tiempo_programado_1)
                    safe_write_cell(ws, 24+row_num_guia, 13, guia.puntaje_1)
                    safe_write_cell(ws, 25+row_num_guia, 14, guia.fecha_entrega_1)
                    safe_write_cell(ws, 31+row_num_guia, 4, guia.tipo_objetivo_2 if hasattr(guia, 'tipo_objetivo_2') else "")
                    safe_write_cell(ws, 31+row_num_guia, 6, guia.objetivo_aprendizaje_2 if hasattr(guia, 'objetivo_aprendizaje_2') else "")
                    safe_write_cell(ws, 31+row_num_guia, 7, guia.contenido_tematico_2 if hasattr(guia, 'contenido_tematico_2') else "")
                    safe_write_cell(ws, 31+row_num_guia, 9, guia.actividad_aprendizaje_2)
                    safe_write_cell(ws, 33+row_num_guia, 9, guia.tecnica_evaluacion_2)
                    safe_write_cell(ws, 34+row_num_guia, 9, guia.tipo_evaluacion_2)
                    safe_write_cell(ws, 35+row_num_guia, 9, guia.instrumento_evaluacion_2)
                    safe_write_cell(ws, 36+row_num_guia, 9, guia.criterios_evaluacion_2)
                    safe_write_cell(ws, 37+row_num_guia, 9, ", ".join(guia.agente_evaluador_2[:2] if guia.agente_evaluador_2 else []))
                    safe_write_cell(ws, 31+row_num_guia, 10, guia.tiempo_minutos_2)
                    safe_write_cell(ws, 31+row_num_guia, 11, guia.recursos_didacticos_2)
                    safe_write_cell(ws, 31+row_num_guia, 13, guia.periodo_tiempo_programado_2)
                    safe_write_cell(ws, 32+row_num_guia, 13, guia.puntaje_2)
                    safe_write_cell(ws, 33+row_num_guia, 14, guia.fecha_entrega_2)
                    safe_write_cell(ws, 39+row_num_guia, 4, guia.tipo_objetivo_3 if hasattr(guia, 'tipo_objetivo_3') else "")
                    safe_write_cell(ws, 39+row_num_guia, 6, guia.objetivo_aprendizaje_3 if hasattr(guia, 'objetivo_aprendizaje_3') else "")
                    safe_write_cell(ws, 39+row_num_guia, 7, guia.contenido_tematico_3 if hasattr(guia, 'contenido_tematico_3') else "")
                    safe_write_cell(ws, 39+row_num_guia, 9, guia.actividad_aprendizaje_3)
                    safe_write_cell(ws, 41+row_num_guia, 9, guia.tecnica_evaluacion_3)
                    safe_write_cell(ws, 42+row_num_guia, 9, guia.tipo_evaluacion_3)
                    safe_write_cell(ws, 43+row_num_guia, 9, guia.instrumento_evaluacion_3)
                    safe_write_cell(ws, 44+row_num_guia, 9, guia.criterios_evaluacion_3)
                    safe_write_cell(ws, 45+row_num_guia, 9, ", ".join(guia.agente_evaluador_3[:2] if guia.agente_evaluador_3 else []))
                    safe_write_cell(ws, 39+row_num_guia, 10, guia.tiempo_minutos_3)
                    safe_write_cell(ws, 39+row_num_guia, 11, guia.recursos_didacticos_3)
                    safe_write_cell(ws, 39+row_num_guia, 13, guia.periodo_tiempo_programado_3)
                    safe_write_cell(ws, 40+row_num_guia, 13, guia.puntaje_3)
                    safe_write_cell(ws, 41+row_num_guia, 14, guia.fecha_entrega_3)
                    safe_write_cell(ws, 47+row_num_guia, 4, guia.tipo_objetivo_4 if hasattr(guia, 'tipo_objetivo_4') else "") # Corrección: Usar tipo_objetivo_4
                    safe_write_cell(ws, 47+row_num_guia, 6, guia.objetivo_aprendizaje_4 if hasattr(guia, 'objetivo_aprendizaje_4') else "") # Corrección: Usar objetivo_aprendizaje_4
                    safe_write_cell(ws, 47+row_num_guia, 7, guia.contenido_tematico_4 if hasattr(guia, 'contenido_tematico_4') else "") # Corrección: Usar contenido_tematico_4
                    safe_write_cell(ws, 47+row_num_guia, 9, guia.actividad_aprendizaje_4)
                    safe_write_cell(ws, 49+row_num_guia, 9, guia.tecnica_evaluacion_4)
                    safe_write_cell(ws, 50+row_num_guia, 9, guia.tipo_evaluacion_4)
                    safe_write_cell(ws, 51+row_num_guia, 9, guia.instrumento_evaluacion_4)
                    safe_write_cell(ws, 52+row_num_guia, 9, guia.criterios_evaluacion_4)
                    safe_write_cell(ws, 53+row_num_guia, 9, ", ".join(guia.agente_evaluador_4[:2] if guia.agente_evaluador_4 else []))
                    safe_write_cell(ws, 47+row_num_guia, 10, guia.tiempo_minutos_4)
                    safe_write_cell(ws, 47+row_num_guia, 11, guia.recursos_didacticos_4)
                    safe_write_cell(ws, 47+row_num_guia, 13, guia.periodo_tiempo_programado_4)
                    safe_write_cell(ws, 48+row_num_guia, 13, guia.puntaje_4)
                    safe_write_cell(ws, 49+row_num_guia, 14, guia.fecha_entrega_4)

                    print(f"Guía (ID: {guia.id}) procesada correctamente para el sílabo {silabo_actual.encuentros}")
                except Exception as e:
                    print(f"Error al procesar la guía del sílabo {silabo_actual.id}: {str(e)}")
                    # Continuamos con el siguiente sílabo sin interrumpir el proceso
            else:
                print(f"El sílabo {silabo_actual.id} no tiene guía asociada.")


        # --- Guarda el archivo Excel en memoria y genera la respuesta ---
        print("Generando respuesta con el archivo Excel...")
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        # Crear un nombre de archivo descriptivo
        filename = f"asignacion_{asignacion_obj.usuario.username}_{asignacion_obj.plan_de_estudio.asignatura.nombre}.xlsx"
        # Sanitizar nombre de archivo si es necesario (remover caracteres especiales)
        filename = "".join(c for c in filename if c.isalnum() or c in ('_', '-')).rstrip()
        response['Content-Disposition'] = f'attachment; filename={filename}'
        wb.save(response)
        print(f"Archivo Excel '{filename}' generado correctamente")
        return response

    except AsignacionPlanEstudio.DoesNotExist:
        print(f"Error: No se encontró la AsignacionPlanEstudio con ID {asignacion_id}")
        return HttpResponse(f"Error: No se encontró la Asignación con ID {asignacion_id}", status=404)
    except Exception as e:
        print(f"Error general durante la generación del Excel para Asignación ID {asignacion_id}: {str(e)}")
        import traceback
        print(traceback.format_exc())
        # Considera retornar un mensaje de error más amigable para el usuario
        return HttpResponse(f"Error interno al generar el archivo Excel: {str(e)}", status=500)
