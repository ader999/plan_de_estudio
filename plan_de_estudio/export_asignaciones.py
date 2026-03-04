from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse

def generar_excel_asignaciones_multiples(queryset, fecha_inicio):
    wb = Workbook()
    
    # Remove default sheet
    if 'Sheet' in wb.sheetnames:
        wb.remove(wb['Sheet'])
        
    # Evaluate queryset to list
    queryset = list(queryset)
    carreras = set([item.plan_de_estudio.carrera.nombre for item in queryset])
    
    if not carreras:
        wb.create_sheet('Sin Datos')
    
    for carrera_nombre in sorted(carreras):
        # Truncate to 31 chars and replace invalid chars for sheet name just in case
        safe_sheet_name = "".join([c for c in carrera_nombre if c.isalnum() or c in " _-"])[:31]
        ws = wb.create_sheet(title=safe_sheet_name) 
        
        # Title
        ws.merge_cells('A1:C1')
        title_cell = ws['A1']
        title_cell.value = f"Carrera: {carrera_nombre}"
        title_cell.font = Font(size=14, bold=True)
        title_cell.alignment = Alignment(horizontal='center')
        
        ws.merge_cells('A2:C2')
        date_cell = ws['A2']
        date_cell.value = f"Fecha de inicio: {fecha_inicio.strftime('%d/%m/%Y')}"
        date_cell.font = Font(italic=True)
        date_cell.alignment = Alignment(horizontal='center')

        current_row = 4
        
        asignaciones_carrera = [item for item in queryset if item.plan_de_estudio.carrera.nombre == carrera_nombre]
        años_en_carrera = set([item.plan_de_estudio.año for item in asignaciones_carrera])
        
        # Define sorting logic for roman numerals if necessary
        roman_order = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5, 'VI': 6, 'VII': 7, 'VIII': 8}
        
        for año in sorted(años_en_carrera, key=lambda x: roman_order.get(x, 99)):
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=3)
            año_cell = ws.cell(row=current_row, column=1)
            año_cell.value = f"Año: {año}"
            año_cell.font = Font(bold=True)
            año_cell.fill = PatternFill(start_color='DDDDDD', end_color='DDDDDD', fill_type='solid')
            current_row += 1
            
            # Headers
            headers = ['Bloque', 'Asignatura', 'Maestro']
            for col_num, header in enumerate(headers, 1):
                cell = ws.cell(row=current_row, column=col_num)
                cell.value = header
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center')
                cell.border = Border(bottom=Side(style='thin'))
            current_row += 1
            
            asignaciones_año = [item for item in asignaciones_carrera if item.plan_de_estudio.año == año]
            # order by bloque
            asignaciones_año = sorted(asignaciones_año, key=lambda x: str(x.bloque if x.bloque else ""))
            
            for asignacion in asignaciones_año:
                row_data = [
                    asignacion.bloque if asignacion.bloque else "-",
                    asignacion.plan_de_estudio.asignatura.nombre,
                    f"{asignacion.usuario.first_name} {asignacion.usuario.last_name}".strip() or asignacion.usuario.username
                ]
                for col_num, value in enumerate(row_data, 1):
                    cell = ws.cell(row=current_row, column=col_num)
                    cell.value = value
                current_row += 1
            
            current_row += 1 # Space between years

        # Auto-adjust columns
        for col in ws.columns:
            max_length = 0
            column_letter = None
            for cell in col:
                # Skip merged cells which might not have simple column_letter via typical access
                if type(cell).__name__ == 'MergedCell':
                    continue
                if not column_letter:
                    column_letter = cell.column_letter
                try:
                    if cell.value and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            if column_letter:
                adjusted_width = (max_length + 2)
                ws.column_dimensions[column_letter].width = adjusted_width

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="asignaciones.xlsx"'
    wb.save(response)
    return response


def generar_word_asignaciones_multiples(queryset, fecha_inicio):
    document = Document()
    
    title = document.add_heading('Asignaciones de Planes de Estudio', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_p = document.add_paragraph()
    date_p.add_run(f"Fecha de inicio: {fecha_inicio.strftime('%d/%m/%Y')}").italic = True
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Evaluate queryset to list to avoid duplicate db calls and preserve select_related cache
    queryset = list(queryset)
    
    carreras = set([item.plan_de_estudio.carrera.nombre for item in queryset])
    
    if not carreras:
        document.add_paragraph("No se encontraron datos para exportar.")
    
    for carrera_nombre in sorted(carreras):
        document.add_heading(f"Carrera: {carrera_nombre}", level=1)
        
        asignaciones_carrera = [item for item in queryset if item.plan_de_estudio.carrera.nombre == carrera_nombre]
        años_en_carrera = set([item.plan_de_estudio.año for item in asignaciones_carrera])
        roman_order = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5, 'VI': 6, 'VII': 7, 'VIII': 8}
        
        for año in sorted(años_en_carrera, key=lambda x: roman_order.get(x, 99)):
            document.add_heading(f"Año: {año}", level=2)
            
            asignaciones_año = [item for item in asignaciones_carrera if item.plan_de_estudio.año == año]
            asignaciones_año = sorted(asignaciones_año, key=lambda x: str(x.bloque if x.bloque else ""))
            
            table = document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Bloque'
            hdr_cells[1].text = 'Asignatura'
            hdr_cells[2].text = 'Maestro'
            
            for asignacion in asignaciones_año:
                row_cells = table.add_row().cells
                row_cells[0].text = asignacion.bloque if asignacion.bloque else "-"
                row_cells[1].text = asignacion.plan_de_estudio.asignatura.nombre
                row_cells[2].text = f"{asignacion.usuario.first_name} {asignacion.usuario.last_name}".strip() or asignacion.usuario.username

            document.add_paragraph() # Spacing

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="asignaciones.docx"'
    document.save(response)
    return response
