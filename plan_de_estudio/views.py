import json
import os
import re
from dotenv import load_dotenv
import logging

from django.db.models.functions import Lower
from django.http import HttpResponse, request ,HttpResponseNotFound, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.conf import settings
import logging

#librerias para el login
from django.contrib.auth import login, authenticate, logout
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import Silabo, Guia, AsignacionPlanEstudio, Asignatura, Plan_de_estudio

from .forms import SilaboForm
from django.views.decorators.csrf import csrf_exempt

from openpyxl import load_workbook
from django.db.models import Count

from openpyxl.utils.dataframe import dataframe_to_rows
import pandas as pd
from docx import Document
from docx.shared import Pt

# Importaciones para IA
import openai  # Importar openai directamente, sin la clase OpenAI
import google.generativeai as genai
import httpx

# Cargar variables de entorno
load_dotenv()
genai.configure(api_key=os.environ.get("GOOGLE_GENERATIVE_API_KEY"))

@login_required
def detalle_silabo(request):
    # Recupera todos los objetos Silabo con sus guías relacionadas
    silabos = Silabo.objects.all().prefetch_related('guias')
    
    # Agrupar sílabos por código de asignatura
    silabos_agrupados = {}
    
    for silabo in silabos:
        codigo = silabo.asignacion.codigo
        if codigo not in silabos_agrupados:
            silabos_agrupados[codigo] = []
        
        silabos_agrupados[codigo].append(silabo)
    
    # Añadir información adicional a cada grupo
    for codigo, grupo in silabos_agrupados.items():
        # Ordenar por número de encuentro
        grupo.sort(key=lambda s: s.encuentros)
        
        # Añadir nombre de asignatura como atributo del grupo
        if grupo:
            grupo.asignatura = grupo[0].asignacion.nombre
    
    return render(request, 'plan_estudio_template/detalle_silabo.html', {'silabos_agrupados': silabos_agrupados})



@login_required
def inicio(request):
    nombre_de_usuario = request.user.username
    # Ajusta el filtro al campo correspondiente en AsignacionPlanEstudio para el usuario
    asignaciones = AsignacionPlanEstudio.objects.filter(usuario__username=nombre_de_usuario)
    rango = range(1, 13)  # Crear el rango para pasarlo al template

    return render(request, 'inicio.html', {'asignaciones': asignaciones, 'rango': rango, 'usuario': nombre_de_usuario})



def acerca_de(request):
    nombre_de_usuario = request.user.username
    return  render(request, 'acerca_de.html',{'usuario':nombre_de_usuario})


def login_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            messages.success(request, 'Inicio de sesión exitoso.')
            return redirect('inicio')  # Cambia 'profile' por la URL deseada después del inicio de sesión
        else:
            messages.error(request, 'Credenciales incorrectas. Por favor, intenta de nuevo.')
    else:  # Si la solicitud es GET, muestra el formulario
        return render(request, 'login.html')

    return render(request, 'login.html')

def logout_view(request):
    logout(request)
    return redirect('login')  # Cambia 'login' por la URL de tu página de inicio de sesión



@login_required
def plan_estudio(request):
    # Obtiene el usuario autenticado
    usuario_autenticado = request.user
    nombre_de_usuario = request.user.username
    # Filtra los silabos del maestro autenticado
    silabos = Silabo.objects.filter(maestro=usuario_autenticado)

    # Crear un diccionario para agrupar los silabos por código
    silabos_agrupados = {}
    for silabo in silabos:
        codigo = silabo.codigo
        if codigo not in silabos_agrupados:
            silabos_agrupados[codigo] = []
        silabos_agrupados[codigo].append(silabo)

    context = {
        'silabos_agrupados': silabos_agrupados,
        'usuario': nombre_de_usuario
    }

    return render(request, 'plan_estudio.html', context)


@login_required
def Plan_de_clase(request):
    return render(request, 'plan_estudio_template/detalle_plandeclase.html')




@login_required
def generar_excel(request):
    print("Me estoy ejecutando")
    if request.method == 'POST':
        # Obtén el código de sílabo del formulario
        codigo_silabo = request.POST.get('codigoSilabo')


        if codigo_silabo:
            # Si se proporcionó un código de sílabo, realiza la búsqueda
            silabos = Silabo.objects.filter(codigo=codigo_silabo, maestro=request.user)

            if silabos.exists():
                # Si se encontraron sílabos, continúa con la lógica para generar el archivo Excel
                # ...

                # Ruta a la plantilla de Excel en tu proyecto
                template_path = 'excel_templates/plantilla.xlsx'

                # Carga la plantilla de Excel
                wb = load_workbook(template_path)

                # Selecciona una hoja de cálculo (worksheet) si es necesario
                ws = wb.active  # O selecciona una hoja específica

                # Inserta los datos en las celdas correspondientes
                row_num = 11  # Fila en la que se insertarán los datos

                for silabo in silabos:
                    ws.cell(row=5, column=2, value=silabo.maestro.username)
                    ws.cell(row=6, column=2, value=silabo.asignatura.asignatura.nombre)
                    ws.cell(row=7, column=2, value=silabo.asignacion_plan.plan_de_estudio.carrera.nombre)
                    ws.cell(row=5, column=4, value=silabo.asignacion_plan.plan_de_estudio.codigo)
                    ws.cell(row=6, column=4, value=silabo.asignacion_plan.plan_de_estudio.año)
                    ws.cell(row=7, column=4, value=silabo.asignacion_plan.plan_de_estudio.trimestre)
                    ws.cell(row=5, column=6, value=silabo.fecha.year)  # Extraemos el año de la fecha del 
                    ws.cell(row=6, column=6, value=silabo.asignacion_plan.plan_de_estudio.hp + silabo.asignacion_plan.plan_de_estudio.hti)  # Extraemos el mes de la fecha del
                    ws.cell(row=7, column=6, value=silabo.asignacion_plan.plan_de_estudio.pr.nombre if silabo.asignacion_plan.plan_de_estudio.pr else "N/A")
                    ws.cell(row=5, column=8, value=silabo.asignacion_plan.plan_de_estudio.pc.nombre if silabo.asignacion_plan.plan_de_estudio.pc else "N/A")
                    ws.cell(row=6, column=8, value=silabo.asignacion_plan.plan_de_estudio.cr.nombre if silabo.asignacion_plan.plan_de_estudio.cr else "N/A")
                    ws.cell(row=row_num, column=1, value=silabo.encuentros)
                    ws.cell(row=row_num, column=2, value=silabo.fecha)
                    ws.cell(row=row_num, column=3, value=silabo.objetivo_conceptual)
                    ws.cell(row=row_num, column=4, value=silabo.objetivo_procedimental)
                    ws.cell(row=row_num, column=5, value=silabo.objetivo_actitudinal)
                    ws.cell(row=row_num, column=6, value=silabo.momento_didactico_primer)
                    ws.cell(row=row_num, column=7, value=silabo.momento_didactico_segundo)
                    ws.cell(row=row_num, column=8, value=silabo.momento_didactico_tercer)
                    ws.cell(row=row_num, column=9, value=silabo.unidad)
                    ws.cell(row=row_num, column=10, value=silabo.contenido_tematico)
                    ws.cell(row=row_num, column=11, value=silabo.forma_organizativa)
                    ws.cell(row=row_num, column=12, value=silabo.tiempo)
                    ws.cell(row=row_num, column=13, value=silabo.tecnicas_aprendizaje)
                    ws.cell(row=row_num, column=14, value=silabo.descripcion_estrategia)
                    ws.cell(row=row_num, column=15, value=silabo.eje_transversal)
                    ws.cell(row=row_num, column=16, value=silabo.hp)

                    # Verificar si el sílabo tiene una guía asociada
                    if hasattr(silabo, 'guia') and silabo.guia is not None:
                        try:
                            ws.cell(row=row_num+16, column=1, value=silabo.guia.numero_guia)
                            ws.cell(row=row_num+16, column=2, value=silabo.guia.fecha)
                            ws.cell(row=row_num+16, column=3, value=silabo.guia.unidad)
                            #ws.cell(row=row_num+16, column=4, value=silabo.guia.nombre_de_la_unidad)
                            ws.cell(row=row_num+16, column=5, value=silabo.guia.objetivo_procedimental)
                            ws.cell(row=row_num+16, column=6, value=silabo.guia.objetivo_conceptual)
                            ws.cell(row=row_num+16, column=7, value=silabo.guia.objetivo_actitudinal)
                            ws.cell(row=row_num+16, column=8, value=silabo.guia.contenido_tematico)
                            ws.cell(row=row_num+16, column=9, value=silabo.guia.actividades)
                            ws.cell(row=row_num+16, column=10, value=silabo.guia.instrumento_cuaderno)
                            ws.cell(row=row_num+16, column=11, value=silabo.guia.instrumento_organizador)
                            ws.cell(row=row_num+16, column=12, value=silabo.guia.instrumento_diario)
                            ws.cell(row=row_num+16, column=13, value=silabo.guia.instrumento_prueba)
                            ws.cell(row=row_num+16, column=14, value=silabo.guia.criterios_evaluacion)
                            ws.cell(row=row_num+16, column=15, value=silabo.guia.tiempo_minutos)
                            ws.cell(row=row_num+16, column=16, value=silabo.guia.recursos)
                            ws.cell(row=row_num+16, column=17, value=silabo.guia.puntaje)
                            ws.cell(row=row_num+16, column=18, value=silabo.guia.evaluacion_sumativa)
                            ws.cell(row=row_num+16, column=19, value=silabo.guia.fecha_entrega)

                        except Exception as e:
                            # Si hay un error al acceder a los datos de la guía, simplemente continuamos
                            print(f"Error al procesar la guía del sílabo {silabo.id}: {str(e)}")


                    # Agrega los datos para otros campos aquí
                    row_num += 1  # Avanza a la siguiente fila

                # Guarda el archivo Excel en memoria
                response = HttpResponse(content_type='application/ms-excel')
                response['Content-Disposition'] = 'attachment; filename=archivo_generado.xlsx'
                wb.save(response)

                return response

    # Si no se proporcionó un código de sílabo o no se encontraron sílabos, redirige a la página principal
    return redirect('plan_de_estudio')  # Reemplaza 'pagina_principal' con la URL de tu elección




@login_required
def generar_excel_original(request):
    print("Me estoy ejecutando")
    if request.method == 'POST':
        # Obtén el código de sílabo del formulario
        codigo_silabo = request.POST.get('codigoSilabo')


        if codigo_silabo:
            # Si se proporcionó un código de sílabo, realiza la búsqueda
            silabos = Silabo.objects.filter(codigo=codigo_silabo, maestro=request.user)

            if silabos.exists():
                # Si se encontraron sílabos, continúa con la lógica para generar el archivo Excel
                # ...

                # Ruta a la plantilla de Excel en tu proyecto
                template_path = 'excel_templates/plantilla_original.xlsx'

                # Carga la plantilla de Excel
                wb = load_workbook(template_path)

                # Selecciona una hoja de cálculo (worksheet) si es necesario
                ws = wb.active  # O selecciona una hoja específica

                # Inserta los datos en las celdas correspondientes
                row_num = 0  # Fila en la que se insertarán los datos

                for silabo in silabos:

                    ws.cell(row=7, column=3, value=silabo.maestro.username)
                    ws.cell(row=5, column=9, value=silabo.asignatura.asignatura.nombre)
                    ws.cell(row=5, column=4, value=silabo.carrera.nombre)
                    ws.cell(row=5, column=13, value=silabo.asignacion_plan.plan_de_estudio.trimestre)
                    ws.cell(row=7, column=10,
                            value=silabo.asignacion_plan.plan_de_estudio.pr.nombre if silabo.asignacion_plan.plan_de_estudio.pr else "N/A")
                    ws.cell(row=7, column=12,
                            value=silabo.asignacion_plan.plan_de_estudio.pc.nombre if silabo.asignacion_plan.plan_de_estudio.pc else "N/A")
                    ws.cell(row=7, column=13,
                            value=silabo.asignacion_plan.plan_de_estudio.cr.nombre if silabo.asignacion_plan.plan_de_estudio.cr else "N/A")

                    ws.cell(row=12+row_num, column=3, value=silabo.unidad)
                    ws.cell(row=13+ row_num, column=3, value=silabo.unidad)
                    ws.cell(row=18+row_num, column=3, value=silabo.unidad)

                    ws.cell(row=11+row_num, column=4, value=silabo.detalle_unidad)
                    ws.cell(row=13+row_num, column=4, value=silabo.detalle_unidad)
                    ws.cell(row=18+row_num, column=4, value=silabo.detalle_unidad)

                    ws.cell(row=11+row_num, column=6, value=silabo.objetivo_conceptual)
                    ws.cell(row=13+row_num, column=6, value=silabo.objetivo_procedimental)
                    ws.cell(row=18+row_num, column=6, value=silabo.objetivo_actitudinal)
                    ws.cell(row=11+row_num, column=7, value=silabo.contenido_tematico)

                    ws.cell(row=12+row_num, column=9, value=silabo.momento_didactico_primer)
                    ws.cell(row=14+row_num, column=9, value=silabo.momento_didactico_segundo)
                    ws.cell(row=19+row_num, column=9, value=silabo.momento_didactico_tercer)

                    ws.cell(row=12+row_num, column=10, value=silabo.tiempo)


                    #Estudio independiente
                    ws.cell(row=23 + row_num, column=3, value=silabo.unidad)
                    ws.cell(row=23 + row_num, column=6, value=silabo.guia.contenido_tematico)
                    ws.cell(row=23 + row_num, column=8, value=silabo.guia.criterios_evaluacion)
                    ws.cell(row=23 + row_num, column=10, value=silabo.guia.tiempo_minutos)
                    ws.cell(row=23 + row_num, column=11, value=silabo.guia.recursos)
                    ws.cell(row=23 + row_num, column=12, value=silabo.guia.puntaje)
                    ws.cell(row=23 + row_num, column=13, value=silabo.guia.evaluacion_sumativa)
                    ws.cell(row=23 + row_num, column=14, value=silabo.guia.objetivo_conceptual)
                    ws.cell(row=23 + row_num, column=15, value=silabo.guia.objetivo_procedimental)
                    ws.cell(row=23 + row_num, column=16, value=silabo.guia.objetivo_actitudinal)
                    ws.cell(row=23 + row_num, column=17, value=silabo.guia.instrumento_cuaderno)
                    ws.cell(row=23 + row_num, column=18, value=silabo.guia.instrumento_organizador)
                    ws.cell(row=23 + row_num, column=19, value=silabo.guia.instrumento_diario)
                    ws.cell(row=23 + row_num, column=20, value=silabo.guia.instrumento_prueba)
                    ws.cell(row=23 + row_num, column=21, value=silabo.guia.fecha_entrega)

                    # Agrega los datos para otros campos aquí
                    row_num += 23  # Avanza a la siguiente fila

                # Guarda el archivo Excel en memoria
                response = HttpResponse(content_type='application/ms-excel')
                response['Content-Disposition'] = 'attachment; filename=archivo_generado.xlsx'
                wb.save(response)

                return response

    # Si no se proporcionó un código de sílabo o no se encontraron sílabos, redirige a la página principal
    return redirect('plan_de_estudio')  # Reemplaza 'pagina_principal' con la URL de tu elección




@login_required
def generar_docx(request):
    if request.method == 'POST':
        codigo_silabo = request.POST.get('codigoSilabo', '')
        if not codigo_silabo:
            return JsonResponse({'error': 'El código de sílabo no se proporcionó.'}, status=400)

        usuario = request.user
        silabos = Silabo.objects.filter(codigo=codigo_silabo, maestro=usuario)

        if not silabos.exists():
            return JsonResponse({'error': 'No se encontraron sílabos para este usuario con el código especificado.'},
                                status=404)

        try:
            # Crear un documento Word
            document = Document()

            # Título principal del documento
            document.add_heading('Sílabo Generado', level=0)

            for silabo in silabos:
                # Encabezado para cada sílabo
                document.add_heading(f'Sílabo: {silabo.codigo}', level=1)

                # Información general
                document.add_paragraph().add_run('Información General').bold = True
                
                # Crear una tabla para la información general
                table = document.add_table(rows=7, cols=2)
                table.style = 'Table Grid'
                
                # Añadir datos a la tabla
                cells = table.rows[0].cells
                cells[0].text = 'Maestro'
                cells[1].text = silabo.maestro.username
                
                cells = table.rows[1].cells
                cells[0].text = 'Asignatura'
                cells[1].text = silabo.asignatura.asignatura.nombre
                
                cells = table.rows[2].cells
                cells[0].text = 'Carrera'
                cells[1].text = silabo.asignacion_plan.plan_de_estudio.carrera.nombre
                
                cells = table.rows[3].cells
                cells[0].text = 'Código Plan'
                cells[1].text = silabo.asignacion_plan.plan_de_estudio.codigo
                
                cells = table.rows[4].cells
                cells[0].text = 'Año'
                cells[1].text = str(silabo.asignacion_plan.plan_de_estudio.año)
                
                cells = table.rows[5].cells
                cells[0].text = 'Trimestre'
                cells[1].text = silabo.asignacion_plan.plan_de_estudio.trimestre
                
                cells = table.rows[6].cells
                cells[0].text = 'Horas Totales'
                cells[1].text = str(silabo.asignacion_plan.plan_de_estudio.hp + silabo.asignacion_plan.plan_de_estudio.hti)
                
                document.add_paragraph()  # Espacio

                # Detalles del sílabo
                document.add_heading('Detalles del Sílabo', level=2)
                
                # Crear una tabla para los detalles
                table = document.add_table(rows=12, cols=2)
                table.style = 'Table Grid'
                
                cells = table.rows[0].cells
                cells[0].text = 'Encuentros'
                cells[1].text = str(silabo.encuentros)
                
                cells = table.rows[1].cells
                cells[0].text = 'Fecha'
                cells[1].text = str(silabo.fecha)
                
                cells = table.rows[2].cells
                cells[0].text = 'Unidad'
                cells[1].text = silabo.unidad
                
                cells = table.rows[3].cells
                cells[0].text = 'Objetivo Conceptual'
                cells[1].text = silabo.objetivo_conceptual
                
                cells = table.rows[4].cells
                cells[0].text = 'Objetivo Procedimental'
                cells[1].text = silabo.objetivo_procedimental
                
                cells = table.rows[5].cells
                cells[0].text = 'Objetivo Actitudinal'
                cells[1].text = silabo.objetivo_actitudinal
                
                cells = table.rows[6].cells
                cells[0].text = 'Primer Momento Didáctico'
                cells[1].text = silabo.momento_didactico_primer
                
                cells = table.rows[7].cells
                cells[0].text = 'Segundo Momento Didáctico'
                cells[1].text = silabo.momento_didactico_segundo
                
                cells = table.rows[8].cells
                cells[0].text = 'Tercer Momento Didáctico'
                cells[1].text = silabo.momento_didactico_tercer
                
                cells = table.rows[9].cells
                cells[0].text = 'Contenido Temático'
                cells[1].text = silabo.contenido_tematico
                
                cells = table.rows[10].cells
                cells[0].text = 'Forma Organizativa'
                cells[1].text = silabo.forma_organizativa
                
                cells = table.rows[11].cells
                cells[0].text = 'Tiempo'
                cells[1].text = str(silabo.tiempo)
                
                document.add_paragraph()  # Espacio
                
                # Información adicional
                document.add_heading('Información Adicional', level=2)
                
                # Crear una tabla para la información adicional
                table = document.add_table(rows=4, cols=2)
                table.style = 'Table Grid'
                
                cells = table.rows[0].cells
                cells[0].text = 'Técnicas de Aprendizaje'
                cells[1].text = silabo.tecnicas_aprendizaje
                
                cells = table.rows[1].cells
                cells[0].text = 'Descripción Estrategia'
                cells[1].text = silabo.descripcion_estrategia
                
                cells = table.rows[2].cells
                cells[0].text = 'Eje Transversal'
                cells[1].text = silabo.eje_transversal
                
                cells = table.rows[3].cells
                cells[0].text = 'Horas Prácticas'
                cells[1].text = str(silabo.hp)
                
                # Verificar si el sílabo tiene una guía asociada
                if hasattr(silabo, 'guia') and silabo.guia is not None:
                    try:
                        document.add_page_break()
                        document.add_heading(f'Guía de Estudio Independiente: {silabo.codigo}', level=1)
                        
                        # Información general de la guía
                        document.add_heading('Información General de la Guía', level=2)
                        
                        # Crear una tabla para la información general de la guía
                        table = document.add_table(rows=4, cols=2)
                        table.style = 'Table Grid'
                        
                        cells = table.rows[0].cells
                        cells[0].text = 'Número de Guía'
                        cells[1].text = str(silabo.guia.numero_guia)
                        
                        cells = table.rows[1].cells
                        cells[0].text = 'Fecha'
                        cells[1].text = str(silabo.guia.fecha)
                        
                        cells = table.rows[2].cells
                        cells[0].text = 'Unidad'
                        cells[1].text = silabo.guia.unidad
                        
                        cells = table.rows[3].cells
                        cells[0].text = 'Contenido Temático'
                        cells[1].text = silabo.guia.contenido_tematico
                        
                        document.add_paragraph()  # Espacio
                        
                        # Objetivos de la guía
                        document.add_heading('Objetivos de la Guía', level=2)
                        
                        # Crear una tabla para los objetivos
                        table = document.add_table(rows=3, cols=2)
                        table.style = 'Table Grid'
                        
                        cells = table.rows[0].cells
                        cells[0].text = 'Objetivo Conceptual'
                        cells[1].text = silabo.guia.objetivo_conceptual
                        
                        cells = table.rows[1].cells
                        cells[0].text = 'Objetivo Procedimental'
                        cells[1].text = silabo.guia.objetivo_procedimental
                        
                        cells = table.rows[2].cells
                        cells[0].text = 'Objetivo Actitudinal'
                        cells[1].text = silabo.guia.objetivo_actitudinal
                        
                        document.add_paragraph()  # Espacio
                        
                        # Actividades y recursos
                        document.add_heading('Actividades y Recursos', level=2)
                        
                        # Crear una tabla para actividades y recursos
                        table = document.add_table(rows=3, cols=2)
                        table.style = 'Table Grid'
                        
                        cells = table.rows[0].cells
                        cells[0].text = 'Actividades'
                        cells[1].text = silabo.guia.actividades
                        
                        cells = table.rows[1].cells
                        cells[0].text = 'Recursos'
                        cells[1].text = silabo.guia.recursos
                        
                        cells = table.rows[2].cells
                        cells[0].text = 'Tiempo (minutos)'
                        cells[1].text = str(silabo.guia.tiempo_minutos)
                        
                        document.add_paragraph()  # Espacio
                        
                        # Evaluación
                        document.add_heading('Evaluación', level=2)
                        
                        # Crear una tabla para la evaluación
                        table = document.add_table(rows=7, cols=2)
                        table.style = 'Table Grid'
                        
                        cells = table.rows[0].cells
                        cells[0].text = 'Instrumento Cuaderno'
                        cells[1].text = silabo.guia.instrumento_cuaderno
                        
                        cells = table.rows[1].cells
                        cells[0].text = 'Instrumento Organizador'
                        cells[1].text = silabo.guia.instrumento_organizador
                        
                        cells = table.rows[2].cells
                        cells[0].text = 'Instrumento Diario'
                        cells[1].text = silabo.guia.instrumento_diario
                        
                        cells = table.rows[3].cells
                        cells[0].text = 'Instrumento Prueba'
                        cells[1].text = silabo.guia.instrumento_prueba
                        
                        cells = table.rows[4].cells
                        cells[0].text = 'Criterios de Evaluación'
                        cells[1].text = silabo.guia.criterios_evaluacion
                        
                        cells = table.rows[5].cells
                        cells[0].text = 'Puntaje'
                        cells[1].text = str(silabo.guia.puntaje)
                        
                        cells = table.rows[6].cells
                        cells[0].text = 'Evaluación Sumativa'
                        cells[1].text = silabo.guia.evaluacion_sumativa
                        
                        document.add_paragraph()  # Espacio
                        
                        # Fecha de entrega
                        document.add_heading('Fecha de Entrega', level=2)
                        document.add_paragraph(str(silabo.guia.fecha_entrega))
                        
                    except Exception as e:
                        # Si hay un error al acceder a los datos de la guía, agregar un mensaje de error
                        document.add_paragraph(f"Error al procesar la guía del sílabo {silabo.id}: {str(e)}")

                # Separador para cada sílabo
                document.add_page_break()

            # Preparar el archivo para descargar
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="silabo_y_guia_generado.docx"'
            document.save(response)

            return response

        except Exception as e:
            return JsonResponse({'error': f"Ocurrió un error al generar el documento Word: {e}"}, status=500)

    return redirect('plan_de_estudio')


@login_required
def success_view(request):
    return render(request, 'exito.html', {
        'message': '¡Gracias por llenar el silabo! Apreciamos el tiempo que has dedicado a completarlo.',
        'usuario': request.user.username,
    })



def obtener_estudios_independientes(asignacion_id):
    """
    Función para obtener todas las guías de estudio independiente asociadas a una asignación.
    
    Args:
        asignacion_id (int): ID de la asignación de plan de estudio
        
    Returns:
        QuerySet: Conjunto de guías de estudio independiente relacionadas con la asignación
    """
    # Primero obtenemos los sílabos relacionados con esta asignación
    silabos = Silabo.objects.filter(asignacion_plan_id=asignacion_id)
    
    # Luego obtenemos todas las guías asociadas a estos sílabos
    guias = Guia.objects.filter(silabo__in=silabos)
    
    return guias


@login_required
def ver_formulario_silabo_guia(request, asignacion_id=None, id=None):
    """
    Función para mostrar el formulario de sílabo y guía.
    Maneja la operación de mostrar el formulario para crear un sílabo o guía.
    """
    if asignacion_id or id:
        if asignacion_id:
            asignacion = get_object_or_404(AsignacionPlanEstudio, id=asignacion_id)
        else:
            asignacion = get_object_or_404(AsignacionPlanEstudio, id=id)
        nombre_de_usuario = request.user.username
        silabos_creados = asignacion.silabo_set.count()

        form = SilaboForm(initial={
            'codigo': asignacion.plan_de_estudio.codigo,
            'carrera': asignacion.plan_de_estudio.carrera,
            'asignatura': asignacion.plan_de_estudio,  # Asignar el plan de estudio como asignatura
            'maestro': request.user,
            'encuentros': silabos_creados + 1,
        })
        guias = obtener_estudios_independientes(asignacion.id)
        form.fields['guia'].queryset = guias

        asignaturas = Asignatura.objects.all()

        return render(request, 'formulario_principal.html', {
            'form': form,
            'asignacion': asignacion,
            'usuario': nombre_de_usuario,
            'silabos_creados': silabos_creados,
            'encuentro': silabos_creados + 1,
            'asignaturas': asignaturas,
        })
    
    # Caso de error: falta asignacion_id
    return JsonResponse({'error': 'Falta ID de asignación'}, status=400)

@login_required
def guardar_silabo_y_guia(request, asignacion_id=None, id=None):
    """
    Función para guardar sílabos y guías de estudio independiente.
    Maneja dos operaciones:
    1. Guardar un nuevo sílabo (método POST con form data)
    2. Agregar una guía de estudio independiente (método POST con JSON data)
    """
    print(f"guardar_silabo_y_guia llamada con asignacion_id={asignacion_id}, id={id}, método={request.method}")
    
    # Caso 1: Guardar guía de estudio y sílabo (POST con JSON)
    if request.method == 'POST' and request.content_type == 'application/json':
        try:
            # Decodificar los datos JSON
            data = json.loads(request.body)
            print(f"Datos JSON recibidos: {data}")
            
            # Obtener el ID de asignación
            asignacion_id = data.get('asignacion_id')
            if not asignacion_id:
                return JsonResponse({'error': 'No se proporcionó un ID de asignación válido'}, status=400)
                
            print(f"Procesando datos para asignacion_id={asignacion_id}")
            
            # Obtener la asignación
            try:
                asignacion = AsignacionPlanEstudio.objects.get(id=asignacion_id)
                print(f"Asignación encontrada: {asignacion}")
            except AsignacionPlanEstudio.DoesNotExist:
                return JsonResponse({'error': f'No se encontró la asignación con ID {asignacion_id}'}, status=404)
            
            # Procesar los datos del sílabo si están presentes
            silabo_data = data.get('silabo_data', {})
            print(f"Datos del sílabo: {silabo_data}")
            
            # Obtener el silabo actual o crear uno nuevo
            silabo = None
            if id:  # Si se está editando un sílabo existente
                silabo = get_object_or_404(Silabo, id=id)
            # Eliminamos la búsqueda del sílabo existente para siempre crear uno nuevo
            # else:
            #     # Buscar si ya existe un sílabo para esta asignación
            #     silabo = Silabo.objects.filter(asignacion_plan=asignacion).order_by('-encuentros').first()
                
            # Si no existe un sílabo o tenemos datos para crear/actualizar uno
            if not silabo or silabo_data:
                # Si tenemos datos del sílabo, creamos o actualizamos
                if silabo_data:
                    # Crear un formulario con los datos recibidos
                    if silabo and id:  # Solo actualizamos si se proporciona un ID específico
                        # Actualizar el sílabo existente
                        form = SilaboForm(silabo_data, instance=silabo)
                    else:
                        # Crear un nuevo sílabo
                        form = SilaboForm(silabo_data)
                    
                    if form.is_valid():
                        silabo = form.save(commit=False)
                        silabo.asignacion_plan = asignacion
                        silabo.maestro = request.user
                        silabo.save()
                        print(f"Sílabo guardado correctamente: {silabo.id}")
                    else:
                        print(f"Errores en el formulario del sílabo: {form.errors}")
                        return JsonResponse({'error': 'Datos del sílabo inválidos', 'form_errors': form.errors}, status=400)
                elif not silabo:
                    # Si no hay sílabo y no tenemos datos, creamos uno con valores predeterminados
                    try:
                        silabo = Silabo.objects.create(
                            codigo=asignacion.plan_de_estudio.codigo,
                            carrera=asignacion.plan_de_estudio.carrera,
                            asignatura=asignacion.plan_de_estudio,
                            maestro=request.user,
                            encuentros=1,
                            fecha=data.get('fecha'),
                            objetivo_conceptual="Objetivo conceptual generado automáticamente",
                            objetivo_procedimental="Objetivo procedimental generado automáticamente",
                            objetivo_actitudinal="Objetivo actitudinal generado automáticamente",
                            momento_didactico_primer="Primer momento didáctico generado automáticamente",
                            momento_didactico_segundo="Segundo momento didáctico generado automáticamente",
                            momento_didactico_tercer="Tercer momento didáctico generado automáticamente",
                            unidad=data.get('unidad', 'Unidad I'),
                            detalle_unidad="Detalle de unidad generado automáticamente",
                            contenido_tematico=data.get('contenido_tematico', 'Contenido temático generado automáticamente'),
                            forma_organizativa="Conferencia",
                            tiempo="60",
                            tecnicas_aprendizaje="Trabajo colaborativo",
                            descripcion_estrategia="Estrategia generada automáticamente",
                            eje_transversal="Investigación",
                            hp="2",
                            asignacion_plan=asignacion
                        )
                        print(f"Sílabo creado automáticamente: {silabo.id}")
                    except Exception as e:
                        print(f"Error al crear el sílabo: {e}")
                        return JsonResponse({'error': f'Error al crear el sílabo: {str(e)}'}, status=400)
            
            # Convertir el tiempo estimado a un float para el campo tiempo_minutos
            try:
                tiempo_minutos = float(data.get('tiempo_minutos', 60.0))
            except (ValueError, TypeError):
                tiempo_minutos = 60.0  # Valor predeterminado si hay un error
                
            # Convertir el puntaje a float si existe
            try:
                puntaje = float(data.get('puntaje')) if data.get('puntaje') else None
            except (ValueError, TypeError):
                puntaje = None
            
            # Mapear los datos recibidos a los campos del modelo Guia
            try:
                # Primero guardamos el sílabo sin asociarlo a una guía
                silabo.guia = None  # Aseguramos que no haya una guía asociada inicialmente
                silabo.save()
                
                # Ahora creamos la guía con referencia al sílabo
                guia = Guia.objects.create(
                    silabo=silabo,  # Asociamos la guía al sílabo
                    numero_guia=int(data.get('numero_guia', 1)),
                    fecha=data.get('fecha'),
                    unidad=data.get('unidad', silabo.unidad),
                    objetivo_conceptual=data.get('objetivo_conceptual', ''),
                    objetivo_procedimental=data.get('objetivo_procedimental', ''),
                    objetivo_actitudinal=data.get('objetivo_actitudinal', ''),
                    contenido_tematico=data.get('contenido_tematico', ''),
                    actividades=data.get('actividades', ''),
                    instrumento_cuaderno=data.get('instrumento_cuaderno', ''),
                    instrumento_organizador=data.get('instrumento_organizador', ''),
                    instrumento_diario=data.get('instrumento_diario', ''),
                    instrumento_prueba=data.get('instrumento_prueba', ''),
                    criterios_evaluacion=data.get('criterios_evaluacion', ''),
                    tiempo_minutos=tiempo_minutos,
                    recursos=data.get('recursos', ''),
                    puntaje=puntaje,
                    evaluacion_sumativa=data.get('evaluacion_sumativa', ''),
                    fecha_entrega=data.get('fecha_entrega')
                )
                print(f"Guía creada correctamente: {guia.id}")
                
                # Ahora actualizamos el sílabo para referenciar esta guía
                silabo.guia = guia
                silabo.save()
                print(f"Sílabo actualizado con referencia a la guía")
                
                # Actualiza el conteo de sílabos creados
                silabos_creados = Silabo.objects.filter(asignacion_plan=asignacion).count()
                asignacion.silabos_creados = silabos_creados
                asignacion.save()
                
                # Redirigir a la vista de éxito
                response_data = {'success': True, 'redirect_url': '/success_view/'}
                return JsonResponse(response_data)
                
            except Exception as e:
                print(f"Error al crear la guía: {e}")
                return JsonResponse({'error': f'Error al crear la guía: {str(e)}'}, status=400)
        except json.JSONDecodeError as e:
            print(f"Error al decodificar JSON: {e}")
            return JsonResponse({'error': f'Error al decodificar JSON: {str(e)}'}, status=400)
        except Exception as e:
            print(f"Error general: {e}")
            return JsonResponse({'error': f'Error general: {str(e)}'}, status=500)
    
    # Caso 2: Guardar sílabo (POST con form data)
    elif request.method == 'POST':
        if asignacion_id:
            asignacion = get_object_or_404(AsignacionPlanEstudio, id=asignacion_id)
        else:
            asignacion = get_object_or_404(AsignacionPlanEstudio, id=id)
        
        form = SilaboForm(request.POST)
        if form.is_valid():
            # Guardar el sílabo
            silabo = form.save(commit=False)
            silabo.asignacion_plan = asignacion
            silabo.asignatura = asignacion.plan_de_estudio  # Aseguramos que se asigne la asignatura
            silabo.save()

            # Actualiza el conteo de sílabos creados
            silabos_creados = asignacion.silabo_set.count()
            asignacion.silabos_creados = silabos_creados
            asignacion.save()

            # Verificar si también se enviaron datos para la guía
            if 'objetivo_conceptual_guia' in request.POST:
                # Primero guardamos el sílabo sin asociarlo a una guía
                silabo.guia = None
                silabo.save()
                
                # Crear una guía asociada al sílabo
                guia = Guia.objects.create(
                    silabo=silabo,  # Asociamos la guía al sílabo
                    numero_guia=1,  # Valor por defecto o ajustar según necesidad
                    fecha=silabo.fecha,  # Usar la misma fecha del sílabo
                    unidad=silabo.unidad,  # Usar la misma unidad del sílabo
                    objetivo_conceptual=request.POST.get('objetivo_conceptual_guia', ''),
                    objetivo_procedimental=request.POST.get('objetivo_procedimental_guia', ''),
                    objetivo_actitudinal=request.POST.get('objetivo_actitudinal_guia', ''),
                    contenido_tematico=request.POST.get('contenido_tematico_guia', silabo.contenido_tematico),
                    actividades=request.POST.get('actividades', ''),
                    instrumento_cuaderno=request.POST.get('instrumento_cuaderno', ''),
                    instrumento_organizador=request.POST.get('instrumento_organizador', ''),
                    instrumento_diario=request.POST.get('instrumento_diario', ''),
                    instrumento_prueba=request.POST.get('instrumento_prueba', ''),
                    criterios_evaluacion=request.POST.get('criterios_evaluacion', ''),
                    tiempo_minutos=float(request.POST.get('tiempo_minutos', 60.0)),
                    recursos=request.POST.get('recursos', ''),
                    evaluacion_sumativa=request.POST.get('evaluacion_sumativa', ''),
                    fecha_entrega=request.POST.get('fecha_entrega', silabo.fecha)
                )
                
                # Actualizar el sílabo para referenciar esta guía
                silabo.guia = guia
                silabo.save()

            messages.success(request, 'El sílabo ha sido creado correctamente.')
            return redirect('success_view')
        else:
            messages.error(request, 'Hubo un error al crear el sílabo. Por favor, revise los datos.')
            # Redirigir al formulario con los errores
            return ver_formulario_silabo_guia(request, asignacion_id, id)
    
    # Caso de error: método no permitido
    return JsonResponse({'error': 'Método no permitido'}, status=405)


@login_required
def generar_silabo(request):
    """
    Función para usar el modelo de Google
    """
    def usar_modelo_google(prompt_completo, generation_config):
        """
        Usa el modelo de Google para generar una respuesta basada en el prompt dado.

        Args:
            prompt_completo (str): Prompt que contiene las instrucciones y datos.
            generation_config (dict): Configuración para la generación del modelo.

        Returns:
            str: Respuesta generada por el modelo.
        """
        # Cargar la clave API desde .env
        load_dotenv()
        api_key = os.environ.get("GOOGLE_GENERATIVE_API_KEY")
        if not api_key:
            raise ValueError("GOOGLE_GENERATIVE_API_KEY no está configurada en el archivo .env")

        try:
            # Configurar la API
            genai.configure(api_key=api_key)

            # Crear el modelo y sesión de chat
            model = genai.GenerativeModel(
                model_name="gemini-1.5-pro-latest",
                generation_config=generation_config
            )
            chat_session = model.start_chat()

            # Generar respuesta
            response = chat_session.send_message(prompt_completo)
            return response.text.strip()

        except genai.errors.GenerativeAIError as e:
            raise RuntimeError(f"Error en la API de Gemini: {e}")
        except Exception as e:
            raise RuntimeError(f"Error inesperado: {e}")


    def usar_modelo_deepseek(prompt_completo, max_tokens=1024, temperature=0.7):
        """
        Usa el modelo de DeepSeek para generar una respuesta basada en el prompt dado.

        Args:
            prompt_completo (str): Prompt que contiene las instrucciones y datos.
            max_tokens (int): Número máximo de tokens en la respuesta.
            temperature (float): Controla la creatividad (0-1).

        Returns:
            str: Respuesta generada por el modelo.
        """
        # Cargar la clave API desde .env
        load_dotenv()
        api_key = os.environ.get("deepseek_API_KEY")
        if not api_key:
            raise ValueError("deepseek_API_KEY no está configurada en el archivo .env")

        api_url = "https://api.deepseek.com/v1/chat/completions"
        
        # Headers de la solicitud
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # Parámetros del cuerpo de la solicitud
        data = {
            "model": "deepseek-chat",  # Usar el modelo correcto según la documentación
            "messages": [
                {"role": "system", "content": "Asistente para crear sílabo y plan de clases."},
                {"role": "user", "content": prompt_completo}
            ],
            "temperature": temperature,
            "max_tokens": max_tokens
        }
        
        try:
            # Hacer la solicitud POST
            import requests
            response = requests.post(api_url, json=data, headers=headers)
            
            # Verificar si la solicitud fue exitosa
            if response.status_code == 200:
                respuesta = response.json()
                return respuesta['choices'][0]['message']['content']
            else:
                error_message = f"Error en la API de DeepSeek: {response.status_code}, {response.text}"
                logging.error(error_message)
                raise RuntimeError(error_message)
                
        except requests.RequestException as e:
            raise RuntimeError(f"Error de conexión con DeepSeek: {e}")
        except Exception as e:
            raise RuntimeError(f"Error inesperado con DeepSeek: {e}")


    def usar_modelo_openai(prompt_completo):
        """
        Función para interactuar con OpenAI usando el cliente de chat basado en el script compartido.
        """
        try:
            # Limpiar variables de entorno de proxy que podrían estar causando el error
            for key in ["HTTP_PROXY", "http_proxy", "HTTPS_PROXY", "https_proxy"]:
               if key in os.environ:
                    del os.environ[key]
            
            # Configurar API key para OpenAI 0.28
            api_key = os.environ.get("OPENAI_API_KEY")
            if not api_key:
                raise ValueError("OPENAI_API_KEY no está configurada en el archivo .env")
                
            # Configurar la API key directamente (estilo OpenAI 0.28)
            import openai
            openai.api_key = api_key
            
            # Crear el mensaje con el modelo usando la API antigua (0.28)
            completion = openai.ChatCompletion.create(
                model="gpt-4o-mini",  # Especifica el modelo deseado
                messages=[
                    {
                        "role": "system",
                        "content": "Asistente para crear sílabo y plan de clases."
                    },
                    {
                        "role": "user",
                        "content": prompt_completo
                    }
                ]
            )

            # Extraer el contenido del mensaje generado
            messages = completion.choices[0].message.content
            print("-------------------------------------------------------------", messages)
            return messages

        except Exception as e:
            print(f"Error detallado al usar OpenAI: {str(e)}")
            raise RuntimeError(f"Error al generar respuesta con OpenAI: {str(e)}")


    def generar_con_openai(prompt):
        """Función interna para generar texto con OpenAI"""
        try:
            # Limpiar variables de entorno de proxy que podrían estar causando el error
            for key in ["HTTP_PROXY", "http_proxy", "HTTPS_PROXY", "https_proxy"]:
               if key in os.environ:
                    del os.environ[key]
            
            # Configurar API key para OpenAI 0.28
            api_key = os.environ.get("OPENAI_API_KEY")
            if not api_key:
                raise ValueError("OPENAI_API_KEY no está configurada en el archivo .env")
                
            # Configurar la API key directamente (estilo OpenAI 0.28)
            import openai
            openai.api_key = api_key
            
            # Crear el mensaje con el modelo usando la API antigua (0.28)
            completion = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",  # Especifica el modelo deseado
                messages=[
                    {
                        "role": "system",
                        "content": "Asistente para generar guías de estudio."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            )
            
            # Extraer el contenido del mensaje
            return completion.choices[0].message.content
        except Exception as e:
            logging.error(f"Error al generar con OpenAI: {str(e)}")
            return None

    if request.method == 'POST':
        # Obtener los datos del formulario
        encuentro = request.POST.get('encuentro')
        plan = request.POST.get('plan')
        # Obtener el modelo seleccionado del formulario
        modelo_seleccionado = request.POST.get('modelo_select')
        print("Imprimiendo el modelo selecionodao:::::::::::::::::: "+str(modelo_seleccionado))
        # Ruta al archivo de datos de ejemplo
        filepath = os.path.join(settings.BASE_DIR, 'static', 'data', 'datos_ejemplo.txt')

        try:
            with open(filepath, 'r') as file:
                datos_ejemplo = file.read()
        except FileNotFoundError:
            return JsonResponse({'error': f"Error: El archivo '{filepath}' no se encuentra."}, status=400)

        # Crear el prompt completo
        prompt_completo = f"""
            Instrucciones: Crea un sílabo basado en la siguiente información y devuélvelo en formato JSON estructurado.

            Datos de ejemplo (para tu referencia):

            {datos_ejemplo}

            el silabo que vas a crear espara el encuentro: {encuentro} de 12 encuentros
            Plan de estudio: {plan}

            Devuelve los datos como un diccionario JSON con las siguientes claves:
            - objetivo_conceptual
            - objetivo_procedimental
            - objetivo_actitudinal
            - momento_didactico_primer
            - momento_didactico_segundo
            - momento_didactico_tercer
            - unidad
            - detalle_unidad
            - contenido_tematico
            - forma_organizativa
            - tiempo
            - tecnicas_aprendizaje
            - descripcion_estrategia
            - eje_transversal
            - hp
            - estudio_independiente
        """
        print("1111111111111111111111111111111111111111111111111", prompt_completo)

        generation_config = {
            "temperature": 0.7,
            "max_output_tokens": 1524
        }

        try:
            # Usar el modelo seleccionado
            if modelo_seleccionado == 'google':
                silabo_generado = usar_modelo_google(prompt_completo, generation_config)
            elif modelo_seleccionado == 'openai':
                silabo_generado = usar_modelo_openai(prompt_completo)
            elif modelo_seleccionado == 'deepseek':
                silabo_generado = usar_modelo_deepseek(prompt_completo, max_tokens=1524, temperature=0.7)
            else:
                return JsonResponse({'error': 'Modelo no válido seleccionado.'}, status=400)

            # Extraer contenido JSON de la respuesta generada
            try:
                # Agregar logging para diagnóstico
                logging.info(f"Respuesta AI recibida (primeros 500 caracteres): {silabo_generado[:500]}...")
                
                # Buscar el bloque JSON con una expresión regular más robusta
                # Busca tanto bloques de código markdown con JSON como JSON directo
                match = re.search(r'```(?:json)?\s*(\{.*\})\s*```|(\{.*\})', silabo_generado, re.DOTALL)
                
                if not match:
                    error_msg = 'No se encontró un bloque JSON válido en la respuesta.'
                    logging.error(f"{error_msg} Respuesta completa: {silabo_generado}")
                    return JsonResponse({'error': error_msg, 'respuesta_completa': silabo_generado[:1000]}, status=500)

                # Extraer el JSON encontrado - puede estar en grupo 1 (dentro de ```) o grupo 2 (JSON directo)
                if match.group(1):
                    silabo_json_text = match.group(1).strip()
                else:
                    silabo_json_text = match.group(2).strip()
                
                logging.info(f"JSON extraído: {silabo_json_text[:200]}...")
                
                # Limpiar cualquier carácter especial o formato
                silabo_json_text = re.sub(r'```json|```', '', silabo_json_text).strip()
                
                # Convertir el texto a un objeto JSON
                try:
                    silabo_data = json.loads(silabo_json_text)
                except json.JSONDecodeError as json_error:
                    # Intento de reparación de JSON si fallan las comillas
                    try:
                        # Reemplazar comillas simples por dobles si es necesario
                        fixed_json = silabo_json_text.replace("'", '"')
                        silabo_data = json.loads(fixed_json)
                        logging.info("JSON reparado exitosamente reemplazando comillas simples por dobles")
                    except json.JSONDecodeError:
                        # Re-lanzar la excepción original si la reparación no funcionó
                        raise json_error
                        
            except json.JSONDecodeError as e:
                error_msg = f'Error al decodificar JSON: {str(e)}'
                logging.error(f"{error_msg} - JSON texto: {silabo_json_text}")
                return JsonResponse({'error': error_msg}, status=500)
            
            # Devolver el JSON como respuesta
            return JsonResponse({'silabo_data': silabo_data})

        except Exception as e:
            print(e)
            return JsonResponse({'error': f"Error general: {str(e)}"}, status=500)

    return JsonResponse({'error': 'Método no permitido.'}, status=405)




@csrf_exempt
def generar_estudio_independiente(request):
    """
    Vista para generar una guía de estudio utilizando IA.
    """
    def usar_modelo_google(prompt_completo, generation_config):
        """
        Usa el modelo de Google para generar una respuesta basada en el prompt dado.

        Args:
            prompt_completo (str): Prompt que contiene las instrucciones y datos.
            generation_config (dict): Configuración para la generación del modelo.

        Returns:
            str: Respuesta generada por el modelo.
        """
        # Cargar la clave API desde .env
        load_dotenv()
        api_key = os.environ.get("GOOGLE_GENERATIVE_API_KEY")
        if not api_key:
            raise ValueError("GOOGLE_GENERATIVE_API_KEY no está configurada en el archivo .env")

        try:
            # Configurar la API
            genai.configure(api_key=api_key)

            # Crear el modelo y sesión de chat
            model = genai.GenerativeModel(
                model_name="gemini-1.5-pro-latest",
                generation_config=generation_config
            )
            chat_session = model.start_chat()

            # Generar respuesta
            response = chat_session.send_message(prompt_completo)
            return response.text.strip()

        except genai.errors.GenerativeAIError as e:
            raise RuntimeError(f"Error en la API de Gemini: {e}")
        except Exception as e:
            raise RuntimeError(f"Error inesperado: {e}")


    def usar_modelo_deepseek(prompt_completo):
        """
        Usa el modelo de DeepSeek para generar una respuesta basada en el prompt dado.

        Args:
            prompt_completo (str): Prompt que contiene las instrucciones y datos.

        Returns:
            str: Respuesta generada por el modelo.
        """
        # Cargar la clave API desde .env
        load_dotenv()
        api_key = os.environ.get("deepseek_API_KEY")
        if not api_key:
            raise ValueError("deepseek_API_KEY no está configurada en el archivo .env")

        api_url = "https://api.deepseek.com/v1/chat/completions"
        
        # Headers de la solicitud
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # Parámetros del cuerpo de la solicitud
        data = {
            "model": "deepseek-chat",  # Usar el modelo correcto según la documentación
            "messages": [
                {"role": "system", "content": "Asistente para crear sílabo y plan de clases."},
                {"role": "user", "content": prompt_completo}
            ],
            "temperature": 0.7,
            "max_tokens": 1024
        }
        
        try:
            # Hacer la solicitud POST
            import requests
            response = requests.post(api_url, json=data, headers=headers)
            
            # Verificar si la solicitud fue exitosa
            if response.status_code == 200:
                respuesta = response.json()
                return respuesta['choices'][0]['message']['content']
            else:
                error_message = f"Error en la API de DeepSeek: {response.status_code}, {response.text}"
                logging.error(error_message)
                raise RuntimeError(error_message)
                
        except requests.RequestException as e:
            raise RuntimeError(f"Error de conexión con DeepSeek: {e}")
        except Exception as e:
            raise RuntimeError(f"Error inesperado con DeepSeek: {e}")


    def usar_modelo_openai(prompt_completo):
        """
        Función para interactuar con OpenAI usando el cliente de chat basado en el script compartido.
        """
        try:
            # Limpiar variables de entorno de proxy que podrían estar causando el error
            for key in ["HTTP_PROXY", "http_proxy", "HTTPS_PROXY", "https_proxy"]:
               if key in os.environ:
                    del os.environ[key]
            
            # Configurar API key para OpenAI 0.28
            api_key = os.environ.get("OPENAI_API_KEY")
            if not api_key:
                raise ValueError("OPENAI_API_KEY no está configurada en el archivo .env")
                
            # Configurar la API key directamente (estilo OpenAI 0.28)
            import openai
            openai.api_key = api_key
            
            # Crear el mensaje con el modelo usando la API antigua (0.28)
            completion = openai.ChatCompletion.create(
                model="gpt-4o-mini",  # Especifica el modelo deseado
                messages=[
                    {
                        "role": "system",
                        "content": "Asistente para crear sílabo y plan de clases."
                    },
                    {
                        "role": "user",
                        "content": prompt_completo
                    }
                ]
            )

            # Extraer el contenido del mensaje generado
            messages = completion.choices[0].message.content
            print("-------------------------------------------------------------", messages)
            return messages

        except Exception as e:
            print(f"Error detallado al usar OpenAI: {str(e)}")
            raise RuntimeError(f"Error al generar respuesta con OpenAI: {str(e)}")


    def generar_con_openai(prompt):
        """Función interna para generar texto con OpenAI"""
        try:
            # Limpiar variables de entorno de proxy que podrían estar causando el error
            for key in ["HTTP_PROXY", "http_proxy", "HTTPS_PROXY", "https_proxy"]:
               if key in os.environ:
                    del os.environ[key]
            
            # Configurar API key para OpenAI 0.28
            api_key = os.environ.get("OPENAI_API_KEY")
            if not api_key:
                raise ValueError("OPENAI_API_KEY no está configurada en el archivo .env")
                
            # Configurar la API key directamente (estilo OpenAI 0.28)
            import openai
            openai.api_key = api_key
            
            # Crear el mensaje con el modelo usando la API antigua (0.28)
            completion = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",  # Especifica el modelo deseado
                messages=[
                    {
                        "role": "system",
                        "content": "Asistente para generar guías de estudio."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            )
            
            # Extraer el contenido del mensaje
            return completion.choices[0].message.content
        except Exception as e:
            logging.error(f"Error al generar con OpenAI: {str(e)}")
            return None

    if request.method == 'POST':
        asignacion_id = request.POST.get('asignacion_id')
        modelo_seleccionado = request.POST.get('modelo_select', 'google')
        
        print(f"Recibida solicitud para generar guía de estudio. Asignación ID: {asignacion_id}, Modelo: {modelo_seleccionado}")

        # Obtener la asignación y datos relacionados
        asignacion = get_object_or_404(AsignacionPlanEstudio, id=asignacion_id)
        plan_estudio = asignacion.plan_de_estudio
        
        # Buscar guías anteriores para usar como contexto principal
        guias_anteriores = Guia.objects.filter(
            silabo__asignacion_plan=asignacion
        ).order_by('numero_guia')
        
        # Obtener el sílabo actual solo para información complementaria
        silabo_actual = Silabo.objects.filter(
            asignacion_plan=asignacion
        ).order_by('-encuentros').first()
        
        # Construir contexto de guías anteriores
        contexto_guias = ""
        if guias_anteriores.exists():
            for idx, guia in enumerate(guias_anteriores):
                # Procesar actividades que pueden estar en formato JSON o texto
                actividades_str = "No especificado"
                if guia.actividades:
                    try:
                        if guia.actividades.startswith('['):
                            actividades_list = json.loads(guia.actividades)
                            actividades_str = ", ".join(actividades_list)
                        else:
                            actividades_str = guia.actividades
                    except:
                        actividades_str = guia.actividades
                
                # Procesar recursos que pueden estar en formato JSON o texto
                recursos_str = "No especificado"
                if guia.recursos:
                    try:
                        if guia.recursos.startswith('['):
                            recursos_list = json.loads(guia.recursos)
                            recursos_str = ", ".join(recursos_list)
                        else:
                            recursos_str = guia.recursos
                    except:
                        recursos_str = guia.recursos
                
                contexto_guias += f"""
                Guía {idx+1}:
                - Unidad: {guia.unidad}
                - Contenido temático: {guia.contenido_tematico or "No especificado"}
                - Actividades: {actividades_str}
                - Recursos: {recursos_str}
                - Tiempo estimado: {guia.tiempo_minutos or "No especificado"} minutos
                - Puntaje: {guia.puntaje or "No especificado"}
                - Evaluación sumativa: {guia.evaluacion_sumativa or "No especificado"}
                - Objetivos: 
                  * Conceptual: {guia.objetivo_conceptual or "No especificado"}
                  * Procedimental: {guia.objetivo_procedimental or "No especificado"}
                  * Actitudinal: {guia.objetivo_actitudinal or "No especificado"}
                - Instrumentos de evaluación:
                  * Cuaderno: {guia.instrumento_cuaderno or "No especificado"}
                  * Organizador gráfico: {guia.instrumento_organizador or "No especificado"}
                  * Diario de trabajo: {guia.instrumento_diario or "No especificado"}
                  * Prueba escrita: {guia.instrumento_prueba or "No especificado"}
                """
        
        # Determinar el número de la próxima guía
        numero_proxima_guia = guias_anteriores.count() + 1
        
        # Crear el prompt completo basado en la información disponible
        if not silabo_actual:
            # Si no hay sílabo, crear un prompt básico con información del plan de estudio
            prompt_completo = f"""
            Instrucciones: Genera una descripción detallada de una guía de estudio inicial (Guía #{numero_proxima_guia}) basada en la siguiente información.
            
            Datos del Plan de Estudio:
            - Asignatura: {plan_estudio.asignatura.nombre}
            - Carrera: {plan_estudio.carrera.nombre}
            - Código: {plan_estudio.codigo}
            - Horas Prácticas (HP): {plan_estudio.hp}
            - Horas de Trabajo Independiente (HTI): {plan_estudio.hti}
            
            {f"Contexto de guías anteriores:{chr(10)}{contexto_guias}" if contexto_guias else "Esta será la primera guía para esta asignación."}
            """
        else:
            # Si hay sílabo, incluir su información
            prompt_completo = f"""
            Instrucciones: Genera una descripción detallada de una guía de estudio (Guía #{numero_proxima_guia}) basada en la siguiente información.
            
            Datos del Sílabo:
            - Asignatura: {plan_estudio.asignatura.nombre}
            - Carrera: {plan_estudio.carrera.nombre}
            - Código: {plan_estudio.codigo}
            - Encuentro actual: {silabo_actual.encuentros} de 12
            - Horas Prácticas (HP): {plan_estudio.hp}
            - Horas de Trabajo Independiente (HTI): {plan_estudio.hti}
            - Unidad actual: {silabo_actual.unidad}
            - Contenido temático: {silabo_actual.contenido_tematico}
            
            {f"Contexto de guías anteriores:{chr(10)}{contexto_guias}" if contexto_guias else "Esta será la primera guía para esta asignación."}
            """
        
        # Añadir la estructura JSON esperada al prompt
        prompt_completo += """
        Por favor, genera una guía de estudio en formato JSON con la siguiente estructura:
        {
          "descripcion": "Una descripción detallada del contenido temático de la guía",
          "actividades": ["Actividad 1", "Actividad 2", "Actividad 3"],
          "recursos": ["Recurso 1", "Recurso 2", "Recurso 3"],
          "tiempo_estimado": "Tiempo estimado en minutos",
          "criterios_evaluacion": ["Criterio 1", "Criterio 2", "Criterio 3"],
          "puntaje": "Puntaje sugerido para esta guía (valor numérico)",
          "evaluacion_sumativa": "Descripción de la evaluación sumativa para esta guía",
          "objetivo_conceptual": "Descripción detallada del objetivo conceptual",
          "objetivo_procedimental": "Descripción detallada del objetivo procedimental",
          "objetivo_actitudinal": "Descripción detallada del objetivo actitudinal",
          "instrumento_cuaderno": "Descripción de cómo se utilizará el cuaderno del estudiante como instrumento de evaluación",
          "instrumento_organizador": "Descripción de cómo se utilizará el organizador gráfico como instrumento de evaluación",
          "instrumento_diario": "Descripción de cómo se utilizará el diario de trabajo como instrumento de evaluación",
          "instrumento_prueba": "Descripción de cómo se utilizará la prueba escrita como instrumento de evaluación"
        }
        
        Asegúrate de que el JSON sea válido y pueda ser parseado correctamente.
        """
        
        print("Prompt generado, enviando a la API...")
        logging.info("Prompt generado, enviando a la API...")
        
        # Generar respuesta según el modelo seleccionado
        respuesta_ai = None
        
        if modelo_seleccionado == 'openai':
            respuesta_ai = generar_con_openai(prompt_completo)
        elif modelo_seleccionado == 'deepseek':
            try:
                respuesta_ai = usar_modelo_deepseek(prompt_completo)
            except Exception as e:
                error_msg = f'Error al generar con DeepSeek: {str(e)}'
                logging.error(error_msg)
                return JsonResponse({'error': error_msg}, status=500)
        else:  # google por defecto
            try:
                # Configurar el modelo usando el método correcto
                generation_config = {
                    "temperature": 0.7,
                    "top_p": 0.95,
                    "top_k": 40,
                    "max_output_tokens": 1524
                }
                
                # Usar la función adecuada para el modelo correcto
                respuesta_ai = usar_modelo_google(prompt_completo, generation_config)
                
            except Exception as e:
                error_msg = f'Error al generar con Google AI: {str(e)}'
                logging.error(error_msg)
                return JsonResponse({'error': error_msg}, status=500)
        
        if not respuesta_ai:
            error_msg = 'No se pudo generar una respuesta'
            logging.error(error_msg)
            return JsonResponse({'error': error_msg}, status=500)
        
        print("Respuesta generada, procesando...")
        logging.info("Respuesta generada, procesando...")
        logging.debug(f"Respuesta AI (primeros 200 caracteres): {respuesta_ai[:200]}...")
        
        # Procesar la respuesta para extraer el JSON
        try:
            # Buscar el JSON en la respuesta
            json_match = re.search(r'\{.*\}', respuesta_ai, re.DOTALL)
            
            if json_match:
                # Extraer el JSON encontrado
                json_str = json_match.group(0).strip()
                
                # Limpiar el string JSON (eliminar bloques de código markdown)
                json_str = re.sub(r'```json|```', '', json_str).strip()
                
                logging.debug(f"JSON extraído (primeros 200 caracteres): {json_str[:200]}...")
                
                # Convertir el texto a un objeto JSON
                data = json.loads(json_str)
                
                # Asegurarse de que los campos esperados estén presentes
                expected_fields = ['descripcion', 'actividades', 'recursos', 'tiempo_estimado', 'criterios_evaluacion', 'puntaje', 'evaluacion_sumativa', 'objetivo_conceptual', 'objetivo_procedimental', 'objetivo_actitudinal', 'instrumento_cuaderno', 'instrumento_organizador', 'instrumento_diario', 'instrumento_prueba']
                for field in expected_fields:
                    if field not in data:
                        data[field] = "" if field != 'tiempo_estimado' else "60"
                        
                    # Asegurarse de que los campos de lista sean realmente listas
                    if field in ['actividades', 'recursos', 'criterios_evaluacion']:
                        if not isinstance(data[field], list):
                            if data[field]:  # Si no está vacío
                                data[field] = [data[field]]
                            else:
                                data[field] = []
                
                print("Datos procesados correctamente, enviando respuesta")
                logging.info("Datos procesados correctamente, enviando respuesta")
                return JsonResponse({'guia_data': data})
            else:
                error_msg = 'No se encontró un formato JSON válido en la respuesta'
                logging.error(error_msg)
                logging.debug(f"Respuesta completa que no contiene JSON: {respuesta_ai}")
                return JsonResponse({'error': error_msg}, status=500)
                    
        except json.JSONDecodeError as e:
            error_msg = f'Error al decodificar JSON: {str(e)}'
            logging.error(error_msg)
            logging.debug(f"JSON con error: {json_str}")
            return JsonResponse({'error': error_msg}, status=500)
        except Exception as e:
            error_msg = f'Error al procesar la respuesta: {str(e)}'
            logging.error(error_msg)
            return JsonResponse({'error': error_msg}, status=500)
        except Exception as e:
            error_msg = f'Error inesperado: {str(e)}'
            logging.error(error_msg)
            return JsonResponse({'error': error_msg}, status=500)


@login_required
def cargar_guia(request, silabo_id):
    """
    Vista para cargar una guía específica de un sílabo mediante AJAX.
    """
    try:
        # Obtener el sílabo específico
        silabo = Silabo.objects.get(id=silabo_id)
        
        # Información de debugging inicial
        debug_info = {
            'silabo_id': silabo_id,
            'silabo_encuentro': silabo.encuentros,
            'asignatura': silabo.asignatura.asignatura if hasattr(silabo, 'asignatura') and silabo.asignatura else 'No asignada'
        }
        
        print(f"DEBUG INFO SILABO: {debug_info}")
        
        # Verificar cuántas guías hay asociadas a este sílabo
        guias_count = Guia.objects.filter(silabo_id=silabo_id).count()
        print(f"Cantidad de guías para el sílabo {silabo_id}: {guias_count}")
        
        if guias_count == 0:
            return render(request, 'plan_estudio_template/detalle_estudioindependiente.html', {
                'silabo': silabo,
                'mensaje_error': 'Este sílabo no tiene guía asociada.',
                'debug_info': debug_info
            })
        elif guias_count > 1:
            # Si hay múltiples guías, tomamos la específica para este encuentro
            # o la primera en su defecto
            print(f"Atención: Se encontraron {guias_count} guías para el sílabo {silabo_id}")
            guias = Guia.objects.filter(silabo_id=silabo_id).order_by('numero_guia')
            for g in guias:
                print(f"Guía ID: {g.id}, Número: {g.numero_guia}, Silabo: {g.silabo_id}")
            
            # Intentamos hacer coincidir la guía con el número de encuentro del sílabo
            try:
                guia = guias.filter(numero_guia=silabo.encuentros).first()
                if not guia:
                    guia = guias.first()  # Si no hay coincidencia, usamos la primera
            except Exception as e:
                print(f"Error al filtrar por encuentro: {str(e)}")
                guia = guias.first()
        else:
            # Si hay exactamente una guía
            guia = Guia.objects.get(silabo_id=silabo_id)
        
        # Información de debugging completa
        debug_info.update({
            'guia_id': guia.id if guia else None,
            'guia_numero': guia.numero_guia if guia else None,
            'total_guias': guias_count
        })
        
        print(f"DEBUG INFO COMPLETO: {debug_info}")
        
        return render(request, 'plan_estudio_template/detalle_estudioindependiente.html', {
            'silabo': silabo,
            'guia': guia,
            'debug_info': debug_info
        })
            
    except Silabo.DoesNotExist:
        return HttpResponse(f"Sílabo no encontrado (ID: {silabo_id})", status=404)
    except Exception as e:
        error_msg = f"Error al cargar la guía: {str(e)} (Sílabo ID: {silabo_id})"
        print(f"ERROR: {error_msg}")
        return HttpResponse(error_msg, status=500)